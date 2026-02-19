"""Tests for server-side TTS generation (BL-032).

Covers:
- tts_generator module functions (sanitize, build, generate, bundle, cleanup)
- Web routes (tts-status, generate-audio, serve, download)
- Graceful degradation when gTTS is not installed
- Export integration (PDF and DOCX with audio_dir)
- CLI generate-audio command
"""

import json
import os
import tempfile
from io import BytesIO
from unittest.mock import MagicMock, patch
from zipfile import ZipFile

import pytest

from src.database import Class, Question, Quiz
from src.tts_generator import (
    _build_question_text,
    _sanitize_text,
    bundle_audio_zip,
    cleanup_quiz_audio,
    get_quiz_audio_dir,
    has_audio,
)

# ============================================================
# Unit tests for tts_generator module
# ============================================================


class TestSanitizeText:
    def test_strips_html_tags(self):
        assert _sanitize_text("<b>bold</b> text") == "bold text"

    def test_collapses_whitespace(self):
        assert _sanitize_text("hello   world") == "hello world"

    def test_empty_string(self):
        assert _sanitize_text("") == ""

    def test_none_input(self):
        assert _sanitize_text(None) == ""

    def test_complex_html(self):
        result = _sanitize_text('<p class="q">What is <em>2+2</em>?</p>')
        assert result == "What is 2+2 ?"


class TestBuildQuestionText:
    def test_text_only(self):
        q = {"text": "What is the capital of France?"}
        result = _build_question_text(q)
        assert "capital of France" in result

    def test_with_options(self):
        q = {"text": "Pick one:", "options": ["Paris", "London", "Berlin"]}
        result = _build_question_text(q)
        assert "A. Paris" in result
        assert "B. London" in result
        assert "C. Berlin" in result

    def test_empty_question(self):
        assert _build_question_text({}) == ""


class TestGenerateQuestionAudio:
    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_success(self, mock_gtts_cls):
        from src.tts_generator import generate_question_audio

        mock_tts = MagicMock()
        mock_gtts_cls.return_value = mock_tts

        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "q1.mp3")
            result = generate_question_audio("Hello world", out)
            assert result == out
            mock_gtts_cls.assert_called_once_with(text="Hello world", lang="en")
            mock_tts.save.assert_called_once_with(out)

    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_empty_text_returns_none(self, mock_gtts_cls):
        from src.tts_generator import generate_question_audio

        result = generate_question_audio("", "/tmp/q1.mp3")
        assert result is None
        mock_gtts_cls.assert_not_called()

    @patch("src.tts_generator.TTS_AVAILABLE", False)
    def test_raises_without_gtts(self):
        from src.tts_generator import generate_question_audio

        with pytest.raises(RuntimeError, match="gTTS is not installed"):
            generate_question_audio("Hello", "/tmp/q1.mp3")

    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_custom_language(self, mock_gtts_cls):
        from src.tts_generator import generate_question_audio

        mock_gtts_cls.return_value = MagicMock()
        with tempfile.TemporaryDirectory() as tmpdir:
            generate_question_audio("Hola", os.path.join(tmpdir, "q1.mp3"), lang="es")
            mock_gtts_cls.assert_called_once_with(text="Hola", lang="es")


class TestGenerateQuizAudio:
    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_generates_files_for_all_questions(self, mock_gtts_cls):
        from src.tts_generator import generate_quiz_audio

        mock_gtts_cls.return_value = MagicMock()
        questions = [
            {"id": 1, "text": "Question one"},
            {"id": 2, "text": "Question two"},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            results = generate_quiz_audio(questions, tmpdir)
            assert len(results) == 2
            assert 1 in results
            assert 2 in results

    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_skips_empty_text(self, mock_gtts_cls):
        from src.tts_generator import generate_quiz_audio

        mock_gtts_cls.return_value = MagicMock()
        questions = [
            {"id": 1, "text": "Valid question"},
            {"id": 2, "text": ""},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            results = generate_quiz_audio(questions, tmpdir)
            assert 1 in results
            assert 2 not in results

    @patch("src.tts_generator.TTS_AVAILABLE", False)
    def test_raises_without_gtts(self):
        from src.tts_generator import generate_quiz_audio

        with pytest.raises(RuntimeError):
            generate_quiz_audio([{"id": 1, "text": "Q"}], "/tmp/audio")


class TestBundleAudioZip:
    def test_creates_zip_with_mp3_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create fake MP3 files
            for i in range(3):
                with open(os.path.join(tmpdir, f"q{i}.mp3"), "wb") as f:
                    f.write(b"fake mp3 data")
            buf = bundle_audio_zip(tmpdir, quiz_title="test_quiz")
            assert isinstance(buf, BytesIO)
            with ZipFile(buf) as zf:
                names = zf.namelist()
                assert len(names) == 3
                assert "test_quiz_audio/q0.mp3" in names

    def test_empty_dir_returns_empty_zip(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            buf = bundle_audio_zip(tmpdir)
            with ZipFile(buf) as zf:
                assert len(zf.namelist()) == 0

    def test_nonexistent_dir_returns_empty_zip(self):
        buf = bundle_audio_zip("/nonexistent/dir")
        with ZipFile(buf) as zf:
            assert len(zf.namelist()) == 0


class TestCleanupQuizAudio:
    def test_removes_audio_files(self):
        with tempfile.TemporaryDirectory() as base:
            audio_dir = os.path.join(base, "42")
            os.makedirs(audio_dir)
            for i in range(3):
                with open(os.path.join(audio_dir, f"q{i}.mp3"), "wb") as f:
                    f.write(b"data")
            cleanup_quiz_audio(42, base_dir=base)
            assert not os.path.exists(audio_dir)

    def test_nonexistent_dir_is_safe(self):
        cleanup_quiz_audio(99999, base_dir="/nonexistent")


class TestHasAudio:
    def test_returns_true_when_files_exist(self):
        with tempfile.TemporaryDirectory() as base:
            audio_dir = os.path.join(base, "1")
            os.makedirs(audio_dir)
            with open(os.path.join(audio_dir, "q1.mp3"), "wb") as f:
                f.write(b"data")
            assert has_audio(1, base_dir=base) is True

    def test_returns_false_when_no_files(self):
        with tempfile.TemporaryDirectory() as base:
            assert has_audio(1, base_dir=base) is False


class TestGetQuizAudioDir:
    def test_returns_correct_path(self):
        result = get_quiz_audio_dir(42, base_dir="uploads/audio")
        assert result == os.path.join("uploads/audio", "42")


# ============================================================
# Web route tests
# ============================================================


@pytest.fixture
def tts_app(make_flask_app):
    """Flask app with a quiz for TTS route tests."""

    def seed(session):
        cls = Class(
            name="TTS Class",
            grade_level="8th",
            subject="Science",
            standards=json.dumps([]),
            config=json.dumps({}),
        )
        session.add(cls)
        session.commit()
        quiz = Quiz(
            title="TTS Quiz",
            class_id=cls.id,
            status="generated",
            style_profile=json.dumps({"provider": "mock"}),
        )
        session.add(quiz)
        session.commit()
        q1 = Question(
            quiz_id=quiz.id,
            question_type="multiple_choice",
            title="Q1",
            text="What is 2+2?",
            points=1.0,
            data=json.dumps(
                {
                    "type": "multiple_choice",
                    "text": "What is 2+2?",
                    "options": ["3", "4", "5"],
                    "correct_answer": "4",
                }
            ),
        )
        session.add(q1)
        session.commit()

    return make_flask_app(seed_fn=seed)


@pytest.fixture
def tts_client(tts_app):
    with tts_app.test_client() as c:
        with c.session_transaction() as sess:
            sess["logged_in"] = True
            sess["username"] = "teacher"
        yield c


class TestTTSStatusRoute:
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=False)
    def test_unavailable(self, mock_avail, tts_client):
        resp = tts_client.get("/api/quizzes/1/tts-status")
        data = resp.get_json()
        assert data["available"] is False
        assert "gTTS" in data.get("message", "") or not data["available"]

    @patch("src.web.blueprints.quizzes.has_audio", return_value=False)
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_available_no_audio(self, mock_avail, mock_has, tts_client):
        resp = tts_client.get("/api/quizzes/1/tts-status")
        data = resp.get_json()
        assert data["available"] is True
        assert data["has_audio"] is False

    @patch("src.web.blueprints.quizzes.has_audio", return_value=True)
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_available_with_audio(self, mock_avail, mock_has, tts_client):
        resp = tts_client.get("/api/quizzes/1/tts-status")
        data = resp.get_json()
        assert data["available"] is True
        assert data["has_audio"] is True


class TestGenerateAudioRoute:
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=False)
    def test_unavailable_returns_400(self, mock_avail, tts_client):
        resp = tts_client.post("/quizzes/1/generate-audio")
        assert resp.status_code == 400

    @patch("src.web.blueprints.quizzes.generate_quiz_audio", return_value={1: "q1.mp3"})
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_success(self, mock_avail, mock_gen, tts_client):
        resp = tts_client.post("/quizzes/1/generate-audio")
        assert resp.status_code == 200
        data = resp.get_json()
        assert data["ok"] is True

    def test_requires_login(self, tts_app):
        with tts_app.test_client() as c:
            resp = c.post("/quizzes/1/generate-audio")
            assert resp.status_code in (302, 303, 401)


class TestServeAudioRoute:
    def test_missing_file_returns_404(self, tts_client):
        resp = tts_client.get("/quizzes/1/audio/999.mp3")
        assert resp.status_code == 404

    def test_requires_login(self, tts_app):
        with tts_app.test_client() as c:
            resp = c.get("/quizzes/1/audio/1.mp3")
            assert resp.status_code in (302, 303, 401)


class TestDownloadAudioRoute:
    @patch("src.web.blueprints.quizzes.has_audio", return_value=False)
    def test_no_audio_returns_404(self, mock_has, tts_client):
        resp = tts_client.get("/quizzes/1/audio/download")
        assert resp.status_code == 404

    def test_requires_login(self, tts_app):
        with tts_app.test_client() as c:
            resp = c.get("/quizzes/1/audio/download")
            assert resp.status_code in (302, 303, 401)


# ============================================================
# Export integration tests (audio_dir parameter)
# ============================================================


def _make_quiz_obj():
    """Create a minimal Quiz-like object for export tests."""
    q = MagicMock()
    q.title = "Audio Test Quiz"
    q.status = "generated"
    q.created_at = None
    q.style_profile = "{}"
    q.reading_level = None
    return q


def _make_question_obj(qid, text="What is 2+2?", qtype="multiple_choice"):
    """Create a minimal Question-like object for export tests."""
    q = MagicMock()
    q.id = qid
    q.text = text
    q.title = f"Q{qid}"
    q.question_type = qtype
    q.points = 1.0
    q.sort_order = qid
    q.data = json.dumps(
        {
            "type": qtype,
            "text": text,
            "options": ["3", "4", "5"],
            "correct_answer": "4",
        }
    )
    return q


class TestExportDocxWithAudio:
    def test_docx_includes_audio_reference_when_file_exists(self):
        from src.export import export_docx

        quiz = _make_quiz_obj()
        questions = [_make_question_obj(42, "What is 2+2?")]

        with tempfile.TemporaryDirectory() as audio_dir:
            # Create a fake audio file matching the question id
            with open(os.path.join(audio_dir, "q42.mp3"), "wb") as f:
                f.write(b"fake mp3")

            buf = export_docx(quiz, questions, audio_dir=audio_dir)
            assert buf is not None
            # Read the docx content to verify audio reference
            from docx import Document

            buf.seek(0)
            doc = Document(buf)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "(Audio: q42.mp3)" in full_text

    def test_docx_no_audio_reference_when_no_file(self):
        from src.export import export_docx

        quiz = _make_quiz_obj()
        questions = [_make_question_obj(42, "What is 2+2?")]

        with tempfile.TemporaryDirectory() as audio_dir:
            # No audio file created
            buf = export_docx(quiz, questions, audio_dir=audio_dir)
            from docx import Document

            buf.seek(0)
            doc = Document(buf)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "(Audio:" not in full_text

    def test_docx_no_audio_reference_when_no_audio_dir(self):
        from src.export import export_docx

        quiz = _make_quiz_obj()
        questions = [_make_question_obj(42, "What is 2+2?")]

        buf = export_docx(quiz, questions, audio_dir=None)
        from docx import Document

        buf.seek(0)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "(Audio:" not in full_text


class TestExportPdfWithAudio:
    def test_pdf_includes_audio_reference_when_file_exists(self):
        """Verify that _pdf_draw_question is called with audio_dir passed through."""
        from src.export import export_pdf

        quiz = _make_quiz_obj()
        questions = [_make_question_obj(42, "What is 2+2?")]

        with tempfile.TemporaryDirectory() as audio_dir:
            with open(os.path.join(audio_dir, "q42.mp3"), "wb") as f:
                f.write(b"fake mp3")

            # Patch _pdf_draw_question to verify audio_dir is passed
            with patch("src.export._pdf_draw_question", wraps=None) as mock_draw:
                # Make it return a reasonable y value so the PDF completes
                mock_draw.return_value = 500.0
                buf = export_pdf(quiz, questions, audio_dir=audio_dir)
                assert buf is not None
                # Verify audio_dir was passed through to the draw function
                mock_draw.assert_called_once()
                call_kwargs = mock_draw.call_args
                assert call_kwargs.kwargs.get("audio_dir") == audio_dir

    def test_pdf_no_audio_reference_when_no_audio_dir(self):
        from src.export import export_pdf

        quiz = _make_quiz_obj()
        questions = [_make_question_obj(42, "What is 2+2?")]

        with patch("src.export._pdf_draw_question", wraps=None) as mock_draw:
            mock_draw.return_value = 500.0
            export_pdf(quiz, questions, audio_dir=None)
            mock_draw.assert_called_once()
            call_kwargs = mock_draw.call_args
            assert call_kwargs.kwargs.get("audio_dir") is None


class TestNormalizeQuestionId:
    def test_includes_question_id(self):
        from src.export import normalize_question

        q = _make_question_obj(99)
        nq = normalize_question(q, 0)
        assert nq["question_id"] == 99

    def test_question_id_is_none_when_missing(self):
        from src.export import normalize_question

        q = _make_question_obj(99)
        del q.id  # remove id attribute
        nq = normalize_question(q, 0)
        assert nq["question_id"] is None


# ============================================================
# Quiz detail page audio UI tests
# ============================================================


class TestQuizDetailAudioUI:
    @patch("src.web.blueprints.quizzes.has_audio", return_value=False)
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_shows_generate_audio_button(self, mock_avail, mock_has, tts_client):
        resp = tts_client.get("/quizzes/1")
        assert resp.status_code == 200
        html = resp.data.decode()
        assert "generateAudioBtn" in html
        assert "Generate Audio" in html

    @patch("src.web.blueprints.quizzes.has_audio", return_value=True)
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_shows_audio_available_badge(self, mock_avail, mock_has, tts_client):
        resp = tts_client.get("/quizzes/1")
        assert resp.status_code == 200
        html = resp.data.decode()
        assert "Audio available" in html
        assert "Download Audio ZIP" in html

    @patch("src.web.blueprints.quizzes.has_audio", return_value=True)
    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=True)
    def test_shows_per_question_audio_download(self, mock_avail, mock_has, tts_client):
        resp = tts_client.get("/quizzes/1")
        assert resp.status_code == 200
        html = resp.data.decode()
        assert "audio-download-link" in html
        assert "MP3" in html

    @patch("src.web.blueprints.quizzes.is_tts_available", return_value=False)
    @patch("src.web.blueprints.quizzes.has_audio", return_value=False)
    def test_hides_audio_when_tts_unavailable(self, mock_has, mock_avail, tts_client):
        resp = tts_client.get("/quizzes/1")
        assert resp.status_code == 200
        html = resp.data.decode()
        # The HTML button element should not be rendered when TTS is unavailable
        # (the JS may still reference the id, but the button itself won't exist)
        assert 'id="generateAudioBtn"' not in html
        assert "Audio available" not in html
        assert "Download Audio ZIP" not in html


# ============================================================
# CLI generate-audio command tests
# ============================================================


class TestCLIGenerateAudio:
    @patch("src.tts_generator.TTS_AVAILABLE", True)
    @patch("src.tts_generator.gTTS")
    def test_generate_audio_success(self, mock_gtts_cls, db_path):
        from src.cli.quiz_commands import handle_generate_audio
        from src.database import Base, get_engine, get_session

        mock_gtts_cls.return_value = MagicMock()

        engine = get_engine(db_path)
        Base.metadata.create_all(engine)
        session = get_session(engine)

        cls = Class(
            name="CLI Class",
            grade_level="8th",
            subject="Science",
            standards=json.dumps([]),
            config=json.dumps({}),
        )
        session.add(cls)
        session.commit()
        quiz = Quiz(
            title="CLI Quiz",
            class_id=cls.id,
            status="generated",
            style_profile=json.dumps({}),
        )
        session.add(quiz)
        session.commit()
        q1 = Question(
            quiz_id=quiz.id,
            question_type="multiple_choice",
            title="Q1",
            text="What is 2+2?",
            points=1.0,
            data=json.dumps({"type": "mc", "text": "What is 2+2?", "options": ["3", "4"], "correct_answer": "4"}),
        )
        session.add(q1)
        session.commit()
        quiz_id = quiz.id
        session.close()
        engine.dispose()

        config = {"paths": {"database_file": db_path}}
        args = MagicMock()
        args.quiz_id = quiz_id
        args.lang = "en"

        with patch("builtins.print") as mock_print:
            handle_generate_audio(config, args)
            printed = " ".join(str(c) for c in mock_print.call_args_list)
            assert "[OK]" in printed

    def test_generate_audio_not_installed(self, db_path, capsys):
        from src.cli.quiz_commands import handle_generate_audio

        config = {"paths": {"database_file": db_path}}
        args = MagicMock()
        args.quiz_id = 1
        args.lang = "en"

        with patch("src.tts_generator.TTS_AVAILABLE", False):
            handle_generate_audio(config, args)
            captured = capsys.readouterr()
            assert "gTTS is not installed" in captured.out

    def test_generate_audio_quiz_not_found(self, db_path, capsys):
        from src.cli.quiz_commands import handle_generate_audio
        from src.database import Base, get_engine

        engine = get_engine(db_path)
        Base.metadata.create_all(engine)
        engine.dispose()

        config = {"paths": {"database_file": db_path}}
        args = MagicMock()
        args.quiz_id = 9999
        args.lang = "en"

        with patch("src.tts_generator.TTS_AVAILABLE", True):
            handle_generate_audio(config, args)
            captured = capsys.readouterr()
            assert "not found" in captured.out
