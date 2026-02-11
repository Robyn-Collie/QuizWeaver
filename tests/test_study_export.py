"""
Tests for QuizWeaver study material export (TSV, CSV, PDF, DOCX).

Covers all export formats and material types.
"""

import csv
import io
import json
import os
import tempfile

import pytest
from docx import Document

from src.database import (
    Base, Class, StudySet, StudyCard,
    get_engine, get_session,
)
from src.study_export import (
    export_flashcards_tsv,
    export_flashcards_csv,
    export_study_pdf,
    export_study_docx,
)


# --- Helpers ---

class FakeStudySet:
    def __init__(self, **kw):
        self.id = kw.get("id", 1)
        self.title = kw.get("title", "Test Study Set")
        self.material_type = kw.get("material_type", "flashcard")
        self.status = kw.get("status", "generated")
        self.created_at = None


class FakeCard:
    def __init__(self, **kw):
        self.id = kw.get("id", 1)
        self.study_set_id = kw.get("study_set_id", 1)
        self.card_type = kw.get("card_type", "flashcard")
        self.sort_order = kw.get("sort_order", 0)
        self.front = kw.get("front", "")
        self.back = kw.get("back", "")
        self.data = kw.get("data", "{}")


def _flashcard_set():
    study_set = FakeStudySet(material_type="flashcard", title="Test Flashcards")
    cards = [
        FakeCard(front="What is DNA?", back="Deoxyribonucleic acid",
                 data=json.dumps({"tags": ["biology", "genetics"]}), sort_order=0),
        FakeCard(front="What is RNA?", back="Ribonucleic acid",
                 data=json.dumps({"tags": ["biology"]}), sort_order=1),
        FakeCard(front="What is ATP?", back="Adenosine triphosphate",
                 data=json.dumps({"tags": ["energy"]}), sort_order=2),
    ]
    return study_set, cards


def _vocab_set():
    study_set = FakeStudySet(material_type="vocabulary", title="Science Vocab")
    cards = [
        FakeCard(card_type="term", front="Mitosis", back="Cell division producing two identical cells",
                 data=json.dumps({"example": "Skin cells divide by mitosis.", "part_of_speech": "noun"})),
        FakeCard(card_type="term", front="Meiosis", back="Cell division producing gametes",
                 data=json.dumps({"example": "Egg and sperm are made by meiosis.", "part_of_speech": "noun"})),
    ]
    return study_set, cards


def _study_guide_set():
    study_set = FakeStudySet(material_type="study_guide", title="Biology Guide")
    cards = [
        FakeCard(card_type="section", front="Introduction", back="Biology is the study of life.",
                 data=json.dumps({"key_points": ["Life requires energy", "Cells are basic units"]})),
        FakeCard(card_type="section", front="Cell Theory", back="All living things are made of cells.",
                 data=json.dumps({"key_points": ["Cells come from cells", "Cell is basic unit"]})),
    ]
    return study_set, cards


def _review_sheet_set():
    study_set = FakeStudySet(material_type="review_sheet", title="Quick Review")
    cards = [
        FakeCard(card_type="fact", front="Photosynthesis Equation", back="6CO2 + 6H2O -> C6H12O6 + 6O2",
                 data=json.dumps({"type": "formula"})),
        FakeCard(card_type="fact", front="Key Fact", back="Mitochondria are the powerhouse of the cell.",
                 data=json.dumps({"type": "fact"})),
    ]
    return study_set, cards


# --- TSV Tests ---

class TestTSVExport:
    def test_tsv_no_header(self):
        study_set, cards = _flashcard_set()
        result = export_flashcards_tsv(study_set, cards)
        lines = result.strip().split("\n")
        assert len(lines) == 3
        # Verify no header (first line has content, not "Front")
        assert "Front" not in lines[0]

    def test_tsv_tab_separated(self):
        study_set, cards = _flashcard_set()
        result = export_flashcards_tsv(study_set, cards)
        lines = result.split("\n")
        for line in lines:
            if not line:
                continue
            parts = line.split("\t")
            assert len(parts) == 4  # front, back, tags, image

    def test_tsv_tags_joined(self):
        study_set, cards = _flashcard_set()
        result = export_flashcards_tsv(study_set, cards)
        lines = result.strip().split("\n")
        # First card has tags ["biology", "genetics"]
        parts = lines[0].split("\t")
        assert "biology" in parts[2]
        assert "genetics" in parts[2]

    def test_tsv_anki_compatible(self):
        """Anki expects: front\\tback\\ttags with no header."""
        study_set, cards = _flashcard_set()
        result = export_flashcards_tsv(study_set, cards)
        # Should not start with # or have a header line
        assert not result.startswith("#")
        assert not result.startswith("Front")


# --- CSV Tests ---

class TestCSVExport:
    def test_csv_has_header(self):
        study_set, cards = _flashcard_set()
        result = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(result))
        header = next(reader)
        assert "Front" in header
        assert "Back" in header

    def test_csv_flashcard_rows(self):
        study_set, cards = _flashcard_set()
        result = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 4  # header + 3 cards

    def test_csv_vocabulary_format(self):
        study_set, cards = _vocab_set()
        result = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(result))
        header = next(reader)
        assert "Term" in header
        assert "Definition" in header
        assert "Example" in header

    def test_csv_review_sheet_format(self):
        study_set, cards = _review_sheet_set()
        result = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(result))
        header = next(reader)
        assert "Type" in header


# --- PDF Tests ---

class TestPDFExport:
    def test_pdf_returns_bytes(self):
        study_set, cards = _flashcard_set()
        result = export_study_pdf(study_set, cards)
        assert result is not None
        data = result.read()
        assert len(data) > 0
        assert data[:5] == b"%PDF-"

    def test_pdf_study_guide(self):
        study_set, cards = _study_guide_set()
        result = export_study_pdf(study_set, cards)
        data = result.read()
        assert len(data) > 0

    def test_pdf_vocabulary(self):
        study_set, cards = _vocab_set()
        result = export_study_pdf(study_set, cards)
        data = result.read()
        assert len(data) > 0

    def test_pdf_review_sheet(self):
        study_set, cards = _review_sheet_set()
        result = export_study_pdf(study_set, cards)
        data = result.read()
        assert len(data) > 0


# --- DOCX Tests ---

class TestDOCXExport:
    def test_docx_returns_bytes(self):
        study_set, cards = _flashcard_set()
        result = export_study_docx(study_set, cards)
        assert result is not None
        data = result.read()
        assert len(data) > 0
        # DOCX files are ZIP archives starting with PK
        assert data[:2] == b"PK"

    def test_docx_flashcard_has_table(self):
        study_set, cards = _flashcard_set()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        assert len(doc.tables) >= 1

    def test_docx_study_guide_has_headings(self):
        study_set, cards = _study_guide_set()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 2

    def test_docx_vocabulary_has_table(self):
        study_set, cards = _vocab_set()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        assert len(doc.tables) >= 1
        # Check header row has Term
        header = doc.tables[0].rows[0]
        header_text = [cell.text for cell in header.cells]
        assert "Term" in header_text
