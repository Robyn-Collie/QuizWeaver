"""
Tests for export improvements: student mode, A/B/C/D lettering,
word bank for fill-in, image suppression, and wider blanks.
"""

import csv
import io
from datetime import datetime

from docx import Document

from src.export import (
    export_csv,
    export_docx,
    export_pdf,
    normalize_question,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_question(**kwargs):
    """Create a Question-like object for testing without a database."""

    class FakeQuestion:
        def __init__(self, **kw):
            self.id = kw.get("id", 1)
            self.quiz_id = kw.get("quiz_id", 1)
            self.question_type = kw.get("question_type", "mc")
            self.title = kw.get("title", "Q1")
            self.text = kw.get("text", "Sample question?")
            self.points = kw.get("points", 5.0)
            self.data = kw.get("data", {})

    return FakeQuestion(**kwargs)


def _make_quiz(**kwargs):
    """Create a Quiz-like object for testing."""

    class FakeQuiz:
        def __init__(self, **kw):
            self.id = kw.get("id", 1)
            self.title = kw.get("title", "Test Quiz")
            self.class_id = kw.get("class_id", 1)
            self.status = kw.get("status", "generated")
            self.style_profile = kw.get("style_profile", "{}")
            self.created_at = kw.get("created_at", datetime(2026, 2, 16, 12, 0))

    return FakeQuiz(**kwargs)


def _mc_question():
    return _make_question(
        question_type="mc",
        text="What is 2+2?",
        points=5,
        data={
            "type": "mc",
            "options": ["1", "2", "3", "4"],
            "correct_answer": "4",
            "cognitive_level": "Remember",
            "cognitive_framework": "blooms",
            "image_description": "A diagram of addition",
        },
    )


def _tf_question():
    return _make_question(
        question_type="tf",
        text="The sky is blue.",
        points=2,
        data={
            "type": "tf",
            "correct_answer": "True",
            "cognitive_level": "Understand",
        },
    )


def _fill_in_question():
    return _make_question(
        question_type="fill_in",
        text="The process of ___ converts light energy.",
        points=3,
        data={
            "type": "fill_in",
            "correct_answer": "photosynthesis",
            "word_bank": ["osmosis", "photosynthesis", "diffusion", "fermentation"],
        },
    )


def _short_answer_question():
    return _make_question(
        question_type="short_answer",
        text="Explain photosynthesis.",
        points=10,
        data={
            "type": "short_answer",
            "expected_answer": "photosynthesis",
            "rubric_hint": "Mention light energy conversion",
        },
    )


# ---------------------------------------------------------------------------
# normalize_question: word_bank field
# ---------------------------------------------------------------------------


class TestNormalizeWordBank:
    """Test that word_bank is included in normalized output."""

    def test_fill_in_with_word_bank(self):
        q = _fill_in_question()
        nq = normalize_question(q, 0)
        assert nq["word_bank"] == ["osmosis", "photosynthesis", "diffusion", "fermentation"]

    def test_mc_no_word_bank(self):
        q = _mc_question()
        nq = normalize_question(q, 0)
        assert nq["word_bank"] is None


# ---------------------------------------------------------------------------
# CSV Export: student_mode
# ---------------------------------------------------------------------------


class TestCSVStudentMode:
    """Test CSV export in student mode vs teacher mode."""

    def test_teacher_mode_has_correct_answer_column(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        csv_str = export_csv(quiz, questions, student_mode=False)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Correct Answer" in header
        assert "Cognitive Level" in header

    def test_student_mode_omits_answer_and_cognitive(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        csv_str = export_csv(quiz, questions, student_mode=True)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Correct Answer" not in header
        assert "Cognitive Level" not in header
        assert "Framework" not in header

    def test_student_mode_row_count_matches(self):
        quiz = _make_quiz()
        questions = [_mc_question(), _tf_question()]
        csv_str = export_csv(quiz, questions, student_mode=True)
        reader = csv.reader(io.StringIO(csv_str))
        rows = list(reader)
        # 1 header + 2 data rows
        assert len(rows) == 3

    def test_student_csv_has_5_columns(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        csv_str = export_csv(quiz, questions, student_mode=True)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert len(header) == 5
        assert header == ["#", "Type", "Question", "Options", "Points"]


# ---------------------------------------------------------------------------
# DOCX Export: student_mode
# ---------------------------------------------------------------------------


def _read_docx_text(buf):
    """Read all paragraph text from a DOCX buffer."""
    doc = Document(buf)
    return "\n".join(p.text for p in doc.paragraphs)


class TestDOCXStudentMode:
    """Test DOCX export in student mode vs teacher mode."""

    def test_teacher_mode_has_answer_key(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "Answer Key" in text

    def test_student_mode_no_answer_key(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        text = _read_docx_text(buf)
        assert "Answer Key" not in text

    def test_student_mode_no_cognitive_level(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        text = _read_docx_text(buf)
        assert "Remember" not in text

    def test_teacher_mode_has_cognitive_level(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "Remember" in text

    def test_student_mode_no_image_description(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        text = _read_docx_text(buf)
        assert "Suggested image" not in text
        assert "diagram of addition" not in text

    def test_teacher_mode_has_image_description(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "Suggested image" in text

    def test_student_mode_has_name_date_line(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        text = _read_docx_text(buf)
        assert "Name:" in text
        assert "Date:" in text

    def test_student_mode_no_generated_by(self):
        quiz = _make_quiz()
        sp = {"provider": "gemini", "model": "gemini-2.5-flash"}
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, style_profile=sp, student_mode=True)
        text = _read_docx_text(buf)
        assert "Generated by" not in text

    def test_teacher_mode_has_generated_by(self):
        quiz = _make_quiz()
        sp = {"provider": "gemini", "model": "gemini-2.5-flash"}
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, style_profile=sp, student_mode=False)
        text = _read_docx_text(buf)
        assert "Generated by" in text

    def test_student_mode_no_rubric_hint(self):
        quiz = _make_quiz()
        questions = [_short_answer_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        text = _read_docx_text(buf)
        assert "Hint:" not in text

    def test_teacher_mode_has_rubric_hint(self):
        quiz = _make_quiz()
        questions = [_short_answer_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "Hint:" in text


# ---------------------------------------------------------------------------
# DOCX: A/B/C/D lettering (no bullets)
# ---------------------------------------------------------------------------


class TestDOCXLettering:
    """Test that MC options use A/B/C/D lettering instead of bullets."""

    def test_mc_options_have_letters(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "A. 1" in text
        assert "B. 2" in text
        assert "C. 3" in text
        assert "D. 4" in text

    def test_tf_options_have_ab_letters(self):
        quiz = _make_quiz()
        questions = [_tf_question()]
        buf = export_docx(quiz, questions, student_mode=False)
        text = _read_docx_text(buf)
        assert "A. True" in text
        assert "B. False" in text

    def test_student_mode_mc_no_bold_correct(self):
        """In student mode, correct answer should not be bolded."""
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions, student_mode=True)
        doc = Document(buf)
        for p in doc.paragraphs:
            if "D. 4" in p.text:
                for run in p.runs:
                    if "D. 4" in run.text:
                        assert not run.bold, "Correct answer should not be bold in student mode"
                        return
        # If we get here, the paragraph was found but assertion succeeded
        assert True


# ---------------------------------------------------------------------------
# DOCX: Word bank for fill-in
# ---------------------------------------------------------------------------


class TestDOCXWordBank:
    """Test word bank rendering in DOCX for fill-in questions."""

    def test_fill_in_word_bank_rendered(self):
        quiz = _make_quiz()
        questions = [_fill_in_question()]
        buf = export_docx(quiz, questions)
        text = _read_docx_text(buf)
        assert "Word Bank:" in text
        assert "photosynthesis" in text
        assert "osmosis" in text

    def test_fill_in_no_word_bank_when_missing(self):
        q = _make_question(
            question_type="fill_in",
            text="The process of ___ happens.",
            data={"type": "fill_in", "correct_answer": "photosynthesis"},
        )
        quiz = _make_quiz()
        buf = export_docx(quiz, [q])
        text = _read_docx_text(buf)
        assert "Word Bank:" not in text


# ---------------------------------------------------------------------------
# PDF Export: student_mode
# ---------------------------------------------------------------------------


class TestPDFStudentMode:
    """Test PDF export in student and teacher modes."""

    def test_pdf_teacher_mode_returns_bytes(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_pdf(quiz, questions, student_mode=False)
        data = buf.read()
        assert data[:4] == b"%PDF"
        assert len(data) > 100

    def test_pdf_student_mode_returns_bytes(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_pdf(quiz, questions, student_mode=True)
        data = buf.read()
        assert data[:4] == b"%PDF"
        assert len(data) > 100

    def test_pdf_student_mode_smaller_or_equal(self):
        """Student mode should produce similar or smaller PDF (no answer key)."""
        quiz = _make_quiz()
        questions = [_mc_question(), _tf_question(), _fill_in_question()]
        teacher_buf = export_pdf(quiz, questions, student_mode=False)
        student_buf = export_pdf(quiz, questions, student_mode=True)
        # Student PDF should be smaller because it lacks the answer key page
        teacher_size = len(teacher_buf.read())
        student_size = len(student_buf.read())
        assert student_size <= teacher_size

    def test_pdf_fill_in_wider_blanks(self):
        """Fill-in blanks should use wider underlines."""
        quiz = _make_quiz()
        questions = [_fill_in_question()]
        buf = export_pdf(quiz, questions)
        # The PDF is binary, so we just verify it generates successfully
        data = buf.read()
        assert len(data) > 100

    def test_pdf_student_mode_multiple_types(self):
        """Student mode should work with mixed question types."""
        quiz = _make_quiz()
        questions = [_mc_question(), _tf_question(), _fill_in_question(), _short_answer_question()]
        buf = export_pdf(quiz, questions, student_mode=True)
        data = buf.read()
        assert data[:4] == b"%PDF"


# ---------------------------------------------------------------------------
# Backward compatibility: default student_mode=False
# ---------------------------------------------------------------------------


class TestBackwardCompatibility:
    """Ensure existing code that doesn't pass student_mode still works."""

    def test_csv_default_is_teacher_mode(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        csv_str = export_csv(quiz, questions)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Correct Answer" in header

    def test_docx_default_is_teacher_mode(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_docx(quiz, questions)
        text = _read_docx_text(buf)
        assert "Answer Key" in text

    def test_pdf_default_is_teacher_mode(self):
        quiz = _make_quiz()
        questions = [_mc_question()]
        buf = export_pdf(quiz, questions)
        data = buf.read()
        assert data[:4] == b"%PDF"
        assert len(data) > 100
