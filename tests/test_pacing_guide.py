"""
Tests for the Curriculum Pacing Guides feature.

Covers:
- CRUD for pacing guides and units
- Template-based generation (quarterly, monthly, semester)
- Current unit lookup by week
- Progress calculation from lesson logs
- CSV, PDF, and DOCX export
- Web routes (auth, CSRF, GET/POST flows)
- Validation (week overlap, required fields)
- Migration SQL
"""

import csv
import io
import json
import os
import sqlite3
import tempfile

import pytest

from src.database import Base, Class, get_engine, get_session, init_db
from src.migrations import run_migrations
from src.pacing_guide import (
    ASSESSMENT_TYPES,
    DEFAULT_TOTAL_WEEKS,
    PACING_TEMPLATES,
    PacingGuideUnit,
    add_unit,
    create_pacing_guide,
    delete_pacing_guide,
    delete_unit,
    generate_from_template,
    get_current_unit,
    get_pacing_guide,
    get_progress,
    list_pacing_guides,
    update_pacing_guide,
    update_unit,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def pacing_db():
    """Provide a temp DB with migrations and a sample class for pacing tests."""
    tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
    tmp.close()
    db_path = tmp.name

    run_migrations(db_path, verbose=False)
    engine = get_engine(db_path)
    init_db(engine)
    session = get_session(engine)

    cls = Class(
        name="Life Science",
        grade_level="7th Grade",
        subject="Science",
        standards=json.dumps(["SOL 7.1", "SOL 7.2", "SOL 7.3"]),
        config=json.dumps({}),
    )
    session.add(cls)
    session.commit()

    yield session, db_path, cls

    session.close()
    engine.dispose()
    try:
        os.remove(db_path)
    except OSError:
        pass


# ---------------------------------------------------------------------------
# TestCreatePacingGuide
# ---------------------------------------------------------------------------


class TestCreatePacingGuide:
    """Tests for creating, getting, listing, and deleting pacing guides."""

    def test_create_basic(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Year Plan")
        assert guide.id is not None
        assert guide.title == "Year Plan"
        assert guide.class_id == cls.id
        assert guide.total_weeks == DEFAULT_TOTAL_WEEKS
        assert guide.school_year is None

    def test_create_with_all_fields(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(
            session, cls.id, "Full Plan", school_year="2025-2026", total_weeks=40
        )
        assert guide.school_year == "2025-2026"
        assert guide.total_weeks == 40

    def test_create_empty_title_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        with pytest.raises(ValueError, match="Title is required"):
            create_pacing_guide(session, cls.id, "")

    def test_create_invalid_class_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        with pytest.raises(ValueError, match="not found"):
            create_pacing_guide(session, 9999, "Test")

    def test_create_zero_weeks_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        with pytest.raises(ValueError, match="total_weeks must be at least 1"):
            create_pacing_guide(session, cls.id, "Test", total_weeks=0)

    def test_get_pacing_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Get Test")
        fetched = get_pacing_guide(session, guide.id)
        assert fetched is not None
        assert fetched.title == "Get Test"

    def test_get_nonexistent_returns_none(self, pacing_db):
        session, db_path, cls = pacing_db
        assert get_pacing_guide(session, 9999) is None

    def test_list_pacing_guides(self, pacing_db):
        session, db_path, cls = pacing_db
        create_pacing_guide(session, cls.id, "Guide A")
        create_pacing_guide(session, cls.id, "Guide B")
        guides = list_pacing_guides(session)
        assert len(guides) >= 2

    def test_list_filtered_by_class(self, pacing_db):
        session, db_path, cls = pacing_db
        cls2 = Class(
            name="Math", grade_level="8th", subject="Math",
            standards=json.dumps([]), config=json.dumps({})
        )
        session.add(cls2)
        session.commit()

        create_pacing_guide(session, cls.id, "Science Plan")
        create_pacing_guide(session, cls2.id, "Math Plan")

        science_guides = list_pacing_guides(session, class_id=cls.id)
        math_guides = list_pacing_guides(session, class_id=cls2.id)
        assert all(g.class_id == cls.id for g in science_guides)
        assert all(g.class_id == cls2.id for g in math_guides)

    def test_delete_pacing_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "To Delete")
        assert delete_pacing_guide(session, guide.id) is True
        assert get_pacing_guide(session, guide.id) is None

    def test_delete_nonexistent_returns_false(self, pacing_db):
        session, db_path, cls = pacing_db
        assert delete_pacing_guide(session, 9999) is False

    def test_update_pacing_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Original")
        updated = update_pacing_guide(session, guide.id, title="Updated", school_year="2026-2027")
        assert updated.title == "Updated"
        assert updated.school_year == "2026-2027"

    def test_update_nonexistent_returns_none(self, pacing_db):
        session, db_path, cls = pacing_db
        assert update_pacing_guide(session, 9999, title="X") is None


# ---------------------------------------------------------------------------
# TestAddUnit
# ---------------------------------------------------------------------------


class TestAddUnit:
    """Tests for adding, updating, and deleting units."""

    def test_add_basic_unit(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        unit = add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        assert unit.id is not None
        assert unit.unit_number == 1
        assert unit.start_week == 1
        assert unit.end_week == 4

    def test_add_unit_with_all_fields(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        unit = add_unit(
            session, guide.id, 1, "Cells", 1, 4,
            standards=["SOL 7.1"], topics=["cells", "organelles"],
            assessment_type="quiz", notes="Focus on plant cells"
        )
        assert json.loads(unit.standards) == ["SOL 7.1"]
        assert json.loads(unit.topics) == ["cells", "organelles"]
        assert unit.assessment_type == "quiz"
        assert unit.notes == "Focus on plant cells"

    def test_add_unit_invalid_guide_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        with pytest.raises(ValueError, match="not found"):
            add_unit(session, 9999, 1, "Test", 1, 4)

    def test_add_unit_empty_title_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        with pytest.raises(ValueError, match="title is required"):
            add_unit(session, guide.id, 1, "", 1, 4)

    def test_add_unit_invalid_weeks_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        with pytest.raises(ValueError, match="start_week must be at least 1"):
            add_unit(session, guide.id, 1, "Test", 0, 4)

    def test_add_unit_end_before_start_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        with pytest.raises(ValueError, match="end_week must be >= start_week"):
            add_unit(session, guide.id, 1, "Test", 5, 3)

    def test_add_unit_exceeds_total_weeks_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide", total_weeks=10)
        with pytest.raises(ValueError, match="exceeds total_weeks"):
            add_unit(session, guide.id, 1, "Test", 1, 11)

    def test_add_unit_invalid_assessment_type_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        with pytest.raises(ValueError, match="Invalid assessment_type"):
            add_unit(session, guide.id, 1, "Test", 1, 4, assessment_type="invalid")

    def test_update_unit(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        unit = add_unit(session, guide.id, 1, "Original", 1, 4)
        updated = update_unit(session, unit.id, title="Updated", assessment_type="test")
        assert updated.title == "Updated"
        assert updated.assessment_type == "test"

    def test_update_unit_lists(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        unit = add_unit(session, guide.id, 1, "Test", 1, 4)
        updated = update_unit(
            session, unit.id,
            standards=["SOL 7.1", "SOL 7.2"],
            topics=["photosynthesis"]
        )
        assert json.loads(updated.standards) == ["SOL 7.1", "SOL 7.2"]
        assert json.loads(updated.topics) == ["photosynthesis"]

    def test_update_nonexistent_returns_none(self, pacing_db):
        session, db_path, cls = pacing_db
        assert update_unit(session, 9999, title="X") is None

    def test_delete_unit(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        unit = add_unit(session, guide.id, 1, "To Delete", 1, 4)
        assert delete_unit(session, unit.id) is True
        assert session.query(PacingGuideUnit).filter_by(id=unit.id).first() is None

    def test_delete_nonexistent_unit_returns_false(self, pacing_db):
        session, db_path, cls = pacing_db
        assert delete_unit(session, 9999) is False

    def test_cascade_delete_units_with_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        add_unit(session, guide.id, 2, "Unit 2", 5, 9)
        guide_id = guide.id
        delete_pacing_guide(session, guide_id)
        units = session.query(PacingGuideUnit).filter_by(pacing_guide_id=guide_id).all()
        assert len(units) == 0


# ---------------------------------------------------------------------------
# TestWeekOverlap
# ---------------------------------------------------------------------------


class TestPacingValidation:
    """Tests for week overlap detection and validation edge cases."""

    def test_week_overlap_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        with pytest.raises(ValueError, match="overlaps"):
            add_unit(session, guide.id, 2, "Unit 2", 3, 6)

    def test_adjacent_weeks_no_overlap(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        unit2 = add_unit(session, guide.id, 2, "Unit 2", 5, 9)
        assert unit2.id is not None

    def test_exact_overlap_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 5, 8)
        with pytest.raises(ValueError, match="overlaps"):
            add_unit(session, guide.id, 2, "Unit 2", 5, 8)

    def test_contained_overlap_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 10)
        with pytest.raises(ValueError, match="overlaps"):
            add_unit(session, guide.id, 2, "Unit 2", 3, 7)


# ---------------------------------------------------------------------------
# TestGenerateFromTemplate
# ---------------------------------------------------------------------------


class TestGenerateFromTemplate:
    """Tests for template-based pacing guide generation."""

    def test_quarterly_template(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Quarterly Plan", "quarterly"
        )
        assert guide.id is not None
        assert guide.total_weeks == DEFAULT_TOTAL_WEEKS
        units = session.query(PacingGuideUnit).filter_by(pacing_guide_id=guide.id).all()
        assert len(units) == 8  # 2 per quarter * 4 quarters

    def test_monthly_template(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Monthly Plan", "monthly"
        )
        units = session.query(PacingGuideUnit).filter_by(pacing_guide_id=guide.id).all()
        assert len(units) == 9  # 1 per month * 9 months

    def test_semester_template(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Semester Plan", "semester"
        )
        units = session.query(PacingGuideUnit).filter_by(pacing_guide_id=guide.id).all()
        assert len(units) == 8  # 4 per semester * 2 semesters

    def test_template_with_standards_distributed(self, pacing_db):
        session, db_path, cls = pacing_db
        standards = ["SOL 7.1", "SOL 7.2", "SOL 7.3", "SOL 7.4"]
        guide = generate_from_template(
            session, cls.id, "With Standards", "quarterly",
            standards_list=standards
        )
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).order_by(PacingGuideUnit.unit_number).all()

        # All standards should be distributed across units
        all_stds = []
        for u in units:
            all_stds.extend(json.loads(u.standards))
        assert sorted(all_stds) == sorted(standards)

    def test_template_units_cover_all_weeks(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Coverage Test", "quarterly"
        )
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).order_by(PacingGuideUnit.start_week).all()

        # First unit starts at week 1
        assert units[0].start_week == 1
        # Last unit ends at total_weeks
        assert units[-1].end_week == guide.total_weeks

    def test_invalid_template_raises(self, pacing_db):
        session, db_path, cls = pacing_db
        with pytest.raises(ValueError, match="Unknown template"):
            generate_from_template(session, cls.id, "Test", "nonexistent")

    def test_template_with_school_year(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Year Plan", "quarterly",
            school_year="2025-2026"
        )
        assert guide.school_year == "2025-2026"

    def test_template_assessment_types_set(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = generate_from_template(
            session, cls.id, "Assessment Test", "quarterly"
        )
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).all()
        # At least some units should have assessment types
        types = [u.assessment_type for u in units if u.assessment_type]
        assert len(types) > 0


# ---------------------------------------------------------------------------
# TestGetCurrentUnit
# ---------------------------------------------------------------------------


class TestGetCurrentUnit:
    """Tests for week-based unit lookup."""

    def test_get_current_unit_found(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9)
        add_unit(session, guide.id, 2, "Unit 2", 10, 18)

        current = get_current_unit(session, guide.id, current_week=5)
        assert current is not None
        assert current.unit_number == 1

    def test_get_current_unit_second(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9)
        add_unit(session, guide.id, 2, "Unit 2", 10, 18)

        current = get_current_unit(session, guide.id, current_week=15)
        assert current is not None
        assert current.unit_number == 2

    def test_get_current_unit_boundary(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9)

        # Exactly at boundary
        current = get_current_unit(session, guide.id, current_week=9)
        assert current is not None
        assert current.unit_number == 1

    def test_get_current_unit_gap_returns_none(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        add_unit(session, guide.id, 2, "Unit 2", 10, 18)

        # Week 7 is in the gap
        current = get_current_unit(session, guide.id, current_week=7)
        assert current is None

    def test_get_current_unit_default_week(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9)

        # Default is week 1
        current = get_current_unit(session, guide.id)
        assert current is not None
        assert current.unit_number == 1


# ---------------------------------------------------------------------------
# TestGetProgress
# ---------------------------------------------------------------------------


class TestGetProgress:
    """Tests for progress calculation from lesson logs."""

    def test_progress_empty_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        progress = get_progress(session, guide.id)
        assert progress["total_units"] == 0
        assert progress["percent_complete"] == 0

    def test_progress_no_lessons(self, pacing_db):
        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9, topics=["cells"])
        progress = get_progress(session, guide.id)
        assert progress["total_units"] == 1
        assert progress["covered_units"] == 0
        assert progress["percent_complete"] == 0

    def test_progress_with_matching_lesson(self, pacing_db):
        from src.database import LessonLog

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9, topics=["cells"])

        # Add a lesson log with matching topic
        log = LessonLog(
            class_id=cls.id,
            content="We studied cells today.",
            topics=json.dumps(["cells"]),
            standards_addressed=json.dumps([]),
        )
        session.add(log)
        session.commit()

        progress = get_progress(session, guide.id)
        assert progress["covered_units"] == 1
        assert progress["percent_complete"] == 100

    def test_progress_partial_coverage(self, pacing_db):
        from src.database import LessonLog

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9, topics=["cells"])
        add_unit(session, guide.id, 2, "Unit 2", 10, 18, topics=["genetics"])

        log = LessonLog(
            class_id=cls.id,
            content="We studied cells.",
            topics=json.dumps(["cells"]),
            standards_addressed=json.dumps([]),
        )
        session.add(log)
        session.commit()

        progress = get_progress(session, guide.id)
        assert progress["total_units"] == 2
        assert progress["covered_units"] == 1
        assert progress["percent_complete"] == 50

    def test_progress_standard_matching(self, pacing_db):
        from src.database import LessonLog

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Test Guide")
        add_unit(session, guide.id, 1, "Unit 1", 1, 9, standards=["SOL 7.1"])

        log = LessonLog(
            class_id=cls.id,
            content="Standards lesson.",
            topics=json.dumps([]),
            standards_addressed=json.dumps(["SOL 7.1"]),
        )
        session.add(log)
        session.commit()

        progress = get_progress(session, guide.id)
        assert progress["covered_units"] == 1

    def test_progress_nonexistent_guide(self, pacing_db):
        session, db_path, cls = pacing_db
        progress = get_progress(session, 9999)
        assert progress["total_units"] == 0


# ---------------------------------------------------------------------------
# TestExportPacingCSV
# ---------------------------------------------------------------------------


class TestExportPacingCSV:
    """Tests for CSV export."""

    def test_csv_format(self, pacing_db):
        from src.pacing_export import export_pacing_csv

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "CSV Test")
        add_unit(
            session, guide.id, 1, "Cells", 1, 4,
            standards=["SOL 7.1"], topics=["cells"], assessment_type="quiz"
        )
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).order_by(PacingGuideUnit.unit_number).all()

        csv_str = export_pacing_csv(guide, units)
        reader = csv.reader(io.StringIO(csv_str))
        rows = list(reader)

        # Header + 1 data row
        assert len(rows) == 2
        assert rows[0] == ["Unit", "Title", "Weeks", "Standards", "Topics", "Assessment", "Notes"]
        assert rows[1][0] == "1"
        assert rows[1][1] == "Cells"
        assert "1-4" in rows[1][2]
        assert "SOL 7.1" in rows[1][3]

    def test_csv_empty_units(self, pacing_db):
        from src.pacing_export import export_pacing_csv

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Empty CSV")
        csv_str = export_pacing_csv(guide, [])
        reader = csv.reader(io.StringIO(csv_str))
        rows = list(reader)
        assert len(rows) == 1  # Header only


# ---------------------------------------------------------------------------
# TestExportPacingPDF
# ---------------------------------------------------------------------------


class TestExportPacingPDF:
    """Tests for PDF export."""

    def test_pdf_returns_bytesio(self, pacing_db):
        from src.pacing_export import export_pacing_pdf

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "PDF Test")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).all()

        buf = export_pacing_pdf(guide, units)
        assert isinstance(buf, io.BytesIO)
        content = buf.read()
        assert len(content) > 0
        assert content[:5] == b"%PDF-"

    def test_pdf_empty_units(self, pacing_db):
        from src.pacing_export import export_pacing_pdf

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "Empty PDF")
        buf = export_pacing_pdf(guide, [])
        assert isinstance(buf, io.BytesIO)
        assert buf.read()[:5] == b"%PDF-"


# ---------------------------------------------------------------------------
# TestExportPacingDOCX
# ---------------------------------------------------------------------------


class TestExportPacingDOCX:
    """Tests for DOCX export."""

    def test_docx_returns_bytesio(self, pacing_db):
        from src.pacing_export import export_pacing_docx

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "DOCX Test")
        add_unit(
            session, guide.id, 1, "Unit 1", 1, 4,
            standards=["SOL 7.1"], topics=["cells"]
        )
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).all()

        buf = export_pacing_docx(guide, units)
        assert isinstance(buf, io.BytesIO)
        content = buf.read()
        assert len(content) > 0
        # DOCX files start with PK (ZIP header)
        assert content[:2] == b"PK"

    def test_docx_table_has_correct_rows(self, pacing_db):
        from docx import Document as DocxDocument

        from src.pacing_export import export_pacing_docx

        session, db_path, cls = pacing_db
        guide = create_pacing_guide(session, cls.id, "DOCX Table Test")
        add_unit(session, guide.id, 1, "Unit 1", 1, 4)
        add_unit(session, guide.id, 2, "Unit 2", 5, 9)
        units = session.query(PacingGuideUnit).filter_by(
            pacing_guide_id=guide.id
        ).order_by(PacingGuideUnit.unit_number).all()

        buf = export_pacing_docx(guide, units)
        doc = DocxDocument(buf)

        # Should have exactly one table
        assert len(doc.tables) == 1
        table = doc.tables[0]
        # Header + 2 data rows
        assert len(table.rows) == 3


# ---------------------------------------------------------------------------
# TestPacingWebRoutes
# ---------------------------------------------------------------------------


class TestPacingWebRoutes:
    """Tests for all pacing guide web routes."""

    @pytest.fixture
    def pacing_client(self, db_path):
        """Provide a logged-in Flask test client with a class seeded."""
        from src.web.app import create_app

        engine = get_engine(db_path)
        Base.metadata.create_all(engine)
        session = get_session(engine)

        cls = Class(
            name="Test Class",
            grade_level="7th Grade",
            subject="Science",
            standards=json.dumps(["SOL 7.1"]),
            config=json.dumps({}),
        )
        session.add(cls)
        session.commit()
        class_id = cls.id
        session.close()
        engine.dispose()

        config = {
            "paths": {"database_file": db_path},
            "llm": {"provider": "mock"},
            "generation": {
                "default_grade_level": "7th Grade",
                "quiz_title": "Test",
                "sol_standards": [],
                "target_image_ratio": 0.0,
                "generate_ai_images": False,
                "interactive_review": False,
            },
        }
        app = create_app(config)
        app.config["TESTING"] = True
        app.config["WTF_CSRF_ENABLED"] = False

        with app.test_client() as client:
            with client.session_transaction() as sess:
                sess["logged_in"] = True
                sess["username"] = "teacher"
            yield client, class_id

        app.config["DB_ENGINE"].dispose()

    def test_list_page(self, pacing_client):
        client, class_id = pacing_client
        resp = client.get("/pacing-guides")
        assert resp.status_code == 200
        assert b"Pacing Guides" in resp.data

    def test_new_page_get(self, pacing_client):
        client, class_id = pacing_client
        resp = client.get("/pacing-guides/new")
        assert resp.status_code == 200
        assert b"Create Pacing Guide" in resp.data

    def test_new_page_post_blank(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/new", data={
            "title": "My Guide",
            "class_id": class_id,
            "total_weeks": "36",
        })
        assert resp.status_code == 303

    def test_new_page_post_with_template(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/new", data={
            "title": "Quarterly Guide",
            "class_id": class_id,
            "total_weeks": "36",
            "template_name": "quarterly",
            "standards_input": "SOL 7.1, SOL 7.2",
        })
        assert resp.status_code == 303

    def test_new_page_post_missing_title(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/new", data={
            "title": "",
            "class_id": class_id,
        })
        assert resp.status_code == 400

    def test_new_page_post_missing_class(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/new", data={
            "title": "Test",
            "class_id": "",
        })
        assert resp.status_code == 400

    def test_detail_page(self, pacing_client):
        client, class_id = pacing_client
        # Create a guide first
        client.post("/pacing-guides/new", data={
            "title": "Detail Test",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1")
        assert resp.status_code == 200
        assert b"Detail Test" in resp.data

    def test_detail_404(self, pacing_client):
        client, class_id = pacing_client
        resp = client.get("/pacing-guides/9999")
        assert resp.status_code == 404

    def test_edit_page_get(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Edit Test",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1/edit")
        assert resp.status_code == 200
        assert b"Edit Pacing Guide" in resp.data

    def test_edit_page_post(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Edit Test",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.post("/pacing-guides/1/edit", data={
            "title": "Updated Title",
            "total_weeks": "40",
        })
        assert resp.status_code == 303

    def test_delete_guide(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Delete Test",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.post("/pacing-guides/1/delete")
        assert resp.status_code == 303

    def test_add_unit_post(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Unit Test",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.post("/pacing-guides/1/add-unit", data={
            "unit_number": "1",
            "title": "Cells",
            "start_week": "1",
            "end_week": "4",
            "standards": "SOL 7.1",
            "topics": "cells, organelles",
            "assessment_type": "quiz",
        })
        assert resp.status_code == 303

    def test_edit_unit_post(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Unit Edit",
            "class_id": class_id,
            "total_weeks": "36",
        })
        client.post("/pacing-guides/1/add-unit", data={
            "unit_number": "1",
            "title": "Original",
            "start_week": "1",
            "end_week": "4",
        })
        resp = client.post("/pacing-guides/1/units/1/edit", data={
            "title": "Updated",
            "unit_number": "1",
            "start_week": "1",
            "end_week": "4",
        })
        assert resp.status_code == 303

    def test_delete_unit_post(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Unit Delete",
            "class_id": class_id,
            "total_weeks": "36",
        })
        client.post("/pacing-guides/1/add-unit", data={
            "unit_number": "1",
            "title": "To Delete",
            "start_week": "1",
            "end_week": "4",
        })
        resp = client.post("/pacing-guides/1/units/1/delete")
        assert resp.status_code == 303

    def test_export_csv(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Export CSV",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1/export/csv")
        assert resp.status_code == 200
        assert resp.content_type.startswith("text/csv")

    def test_export_pdf(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Export PDF",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1/export/pdf")
        assert resp.status_code == 200
        assert resp.content_type == "application/pdf"

    def test_export_docx(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Export DOCX",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1/export/docx")
        assert resp.status_code == 200

    def test_export_invalid_format_404(self, pacing_client):
        client, class_id = pacing_client
        client.post("/pacing-guides/new", data={
            "title": "Export Bad",
            "class_id": class_id,
            "total_weeks": "36",
        })
        resp = client.get("/pacing-guides/1/export/xlsx")
        assert resp.status_code == 404

    def test_generate_post(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/generate", data={
            "title": "Generated",
            "class_id": class_id,
            "template_name": "monthly",
            "school_year": "2025-2026",
        })
        assert resp.status_code == 303

    def test_generate_missing_fields_redirects(self, pacing_client):
        client, class_id = pacing_client
        resp = client.post("/pacing-guides/generate", data={
            "title": "",
            "class_id": class_id,
            "template_name": "",
        })
        assert resp.status_code == 303

    def test_auth_required(self, db_path):
        """Verify unauthenticated users are redirected."""
        from src.web.app import create_app

        config = {
            "paths": {"database_file": db_path},
            "llm": {"provider": "mock"},
            "generation": {
                "default_grade_level": "7th Grade",
                "quiz_title": "Test",
                "sol_standards": [],
                "target_image_ratio": 0.0,
                "generate_ai_images": False,
                "interactive_review": False,
            },
        }
        app = create_app(config)
        app.config["TESTING"] = True
        app.config["WTF_CSRF_ENABLED"] = False

        with app.test_client() as client:
            resp = client.get("/pacing-guides")
            assert resp.status_code == 303
            assert "/login" in resp.headers.get("Location", "")

        app.config["DB_ENGINE"].dispose()


# ---------------------------------------------------------------------------
# TestMigration
# ---------------------------------------------------------------------------


class TestMigration:
    """Tests for the migration SQL creating pacing guide tables."""

    def test_migration_creates_tables(self):
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.close()
        db_path = tmp.name

        try:
            conn = sqlite3.connect(db_path)
            # Read and execute migration SQL
            migration_path = os.path.join(
                os.path.dirname(os.path.dirname(__file__)),
                "migrations",
                "011_pacing_guides.sql",
            )
            with open(migration_path) as f:
                sql = f.read()
            conn.executescript(sql)

            # Verify tables exist
            cursor = conn.cursor()
            cursor.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='pacing_guides'"
            )
            assert cursor.fetchone() is not None

            cursor.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='pacing_guide_units'"
            )
            assert cursor.fetchone() is not None

            # Verify columns on pacing_guides
            cursor.execute("PRAGMA table_info(pacing_guides)")
            pg_columns = {row[1] for row in cursor.fetchall()}
            assert "id" in pg_columns
            assert "class_id" in pg_columns
            assert "title" in pg_columns
            assert "school_year" in pg_columns
            assert "total_weeks" in pg_columns

            # Verify columns on pacing_guide_units
            cursor.execute("PRAGMA table_info(pacing_guide_units)")
            pgu_columns = {row[1] for row in cursor.fetchall()}
            assert "pacing_guide_id" in pgu_columns
            assert "unit_number" in pgu_columns
            assert "start_week" in pgu_columns
            assert "end_week" in pgu_columns
            assert "standards" in pgu_columns
            assert "topics" in pgu_columns
            assert "assessment_type" in pgu_columns

            conn.close()
        finally:
            try:
                os.remove(db_path)
            except OSError:
                pass

    def test_migration_idempotent(self):
        """Running migration twice should not fail."""
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.close()
        db_path = tmp.name

        try:
            conn = sqlite3.connect(db_path)
            migration_path = os.path.join(
                os.path.dirname(os.path.dirname(__file__)),
                "migrations",
                "011_pacing_guides.sql",
            )
            with open(migration_path) as f:
                sql = f.read()
            # Run twice -- should not raise
            conn.executescript(sql)
            conn.executescript(sql)
            conn.close()
        finally:
            try:
                os.remove(db_path)
            except OSError:
                pass


# ---------------------------------------------------------------------------
# TestConstants
# ---------------------------------------------------------------------------


class TestConstants:
    """Tests for module-level constants and templates."""

    def test_default_total_weeks(self):
        assert DEFAULT_TOTAL_WEEKS == 36

    def test_assessment_types(self):
        assert "quiz" in ASSESSMENT_TYPES
        assert "test" in ASSESSMENT_TYPES
        assert "project" in ASSESSMENT_TYPES
        assert "exit_ticket" in ASSESSMENT_TYPES
        assert "performance_task" in ASSESSMENT_TYPES

    def test_pacing_templates_keys(self):
        assert "quarterly" in PACING_TEMPLATES
        assert "monthly" in PACING_TEMPLATES
        assert "semester" in PACING_TEMPLATES

    def test_pacing_templates_have_labels(self):
        for key, tmpl in PACING_TEMPLATES.items():
            assert "label" in tmpl
            assert "description" in tmpl
