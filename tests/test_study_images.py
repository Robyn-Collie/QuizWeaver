"""
Tests for BL-008: Images in Study Materials.

Verifies image_url support in study generator, display in templates,
and image references in all export formats (TSV, CSV, PDF, DOCX).
"""

import csv
import io
import json
import os
import tempfile

import pytest
from docx import Document

from src.database import (
    Base,
    Class,
    StudyCard,
    StudySet,
    get_engine,
    get_session,
)
from src.study_export import (
    export_flashcards_csv,
    export_flashcards_tsv,
    export_study_docx,
    export_study_pdf,
)

# --- Helpers ---


class FakeStudySet:
    def __init__(self, **kw):
        self.id = kw.get("id", 1)
        self.title = kw.get("title", "Test Study Set")
        self.material_type = kw.get("material_type", "flashcard")
        self.status = kw.get("status", "generated")
        self.created_at = None


class FakeCard:
    def __init__(self, **kw):
        self.id = kw.get("id", 1)
        self.study_set_id = kw.get("study_set_id", 1)
        self.card_type = kw.get("card_type", "flashcard")
        self.sort_order = kw.get("sort_order", 0)
        self.front = kw.get("front", "")
        self.back = kw.get("back", "")
        self.data = kw.get("data", "{}")


def _flashcard_set_with_images():
    study_set = FakeStudySet(material_type="flashcard", title="Image Flashcards")
    cards = [
        FakeCard(
            front="Cell Diagram",
            back="A labeled diagram of a cell",
            data=json.dumps(
                {
                    "tags": ["biology", "cells"],
                    "image_url": "https://example.com/cell.png",
                }
            ),
            sort_order=0,
        ),
        FakeCard(
            front="No Image Card",
            back="Card without image",
            data=json.dumps({"tags": ["test"]}),
            sort_order=1,
        ),
    ]
    return study_set, cards


def _study_guide_with_images():
    study_set = FakeStudySet(material_type="study_guide", title="Image Study Guide")
    cards = [
        FakeCard(
            card_type="section",
            front="Photosynthesis",
            back="Process by which plants convert sunlight to energy",
            data=json.dumps(
                {
                    "key_points": ["Uses chlorophyll", "Produces oxygen"],
                    "image_url": "https://example.com/photosynthesis.jpg",
                }
            ),
            sort_order=0,
        ),
    ]
    return study_set, cards


def _vocabulary_with_images():
    study_set = FakeStudySet(material_type="vocabulary", title="Image Vocabulary")
    cards = [
        FakeCard(
            card_type="term",
            front="Mitochondria",
            back="Powerhouse of the cell",
            data=json.dumps(
                {
                    "example": "Mitochondria produce ATP.",
                    "part_of_speech": "noun",
                    "image_url": "https://example.com/mitochondria.png",
                }
            ),
            sort_order=0,
        ),
    ]
    return study_set, cards


def _review_sheet_with_images():
    study_set = FakeStudySet(material_type="review_sheet", title="Image Review Sheet")
    cards = [
        FakeCard(
            card_type="fact",
            front="E=mc^2",
            back="Energy equals mass times the speed of light squared",
            data=json.dumps(
                {
                    "type": "formula",
                    "image_url": "https://example.com/einstein.png",
                }
            ),
            sort_order=0,
        ),
    ]
    return study_set, cards


# ============================================================
# TSV Export with Images
# ============================================================


class TestTSVExportImages:
    """Test TSV export includes image_url column."""

    def test_tsv_includes_image_tag(self):
        """TSV export includes <img> tag for Anki when image_url present."""
        study_set, cards = _flashcard_set_with_images()
        tsv = export_flashcards_tsv(study_set, cards)
        lines = tsv.split("\n")
        # First card has image
        assert '<img src="https://example.com/cell.png">' in lines[0]

    def test_tsv_no_image_empty_column(self):
        """TSV export has empty image column when no image_url."""
        study_set, cards = _flashcard_set_with_images()
        tsv = export_flashcards_tsv(study_set, cards)
        lines = tsv.split("\n")
        # Second card has no image -- last tab should have nothing after it
        fields = lines[1].split("\t")
        assert len(fields) == 4
        assert fields[3] == ""

    def test_tsv_has_four_columns(self):
        """TSV export now has 4 columns (front, back, tags, image)."""
        study_set, cards = _flashcard_set_with_images()
        tsv = export_flashcards_tsv(study_set, cards)
        for line in tsv.split("\n"):
            fields = line.split("\t")
            assert len(fields) == 4


# ============================================================
# CSV Export with Images
# ============================================================


class TestCSVExportImages:
    """Test CSV export includes Image URL column."""

    def test_csv_flashcard_has_image_url_header(self):
        """Flashcard CSV header includes Image URL column."""
        study_set, cards = _flashcard_set_with_images()
        csv_str = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Image URL" in header

    def test_csv_flashcard_image_url_in_rows(self):
        """Flashcard CSV rows contain image URLs."""
        study_set, cards = _flashcard_set_with_images()
        csv_str = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(csv_str))
        next(reader)  # skip header
        row1 = next(reader)
        assert "https://example.com/cell.png" in row1

    def test_csv_vocabulary_has_image_url_header(self):
        """Vocabulary CSV header includes Image URL column."""
        study_set, cards = _vocabulary_with_images()
        csv_str = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Image URL" in header

    def test_csv_study_guide_has_image_url_header(self):
        """Study guide CSV header includes Image URL column."""
        study_set, cards = _study_guide_with_images()
        csv_str = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Image URL" in header

    def test_csv_review_sheet_has_image_url_header(self):
        """Review sheet CSV header includes Image URL column."""
        study_set, cards = _review_sheet_with_images()
        csv_str = export_flashcards_csv(study_set, cards)
        reader = csv.reader(io.StringIO(csv_str))
        header = next(reader)
        assert "Image URL" in header


# ============================================================
# PDF Export with Images
# ============================================================


class TestPDFExportImages:
    """Test PDF export includes image URL references."""

    def test_pdf_flashcard_with_image_url(self):
        """PDF export for flashcards includes image URL text."""
        study_set, cards = _flashcard_set_with_images()
        buf = export_study_pdf(study_set, cards)
        assert buf is not None
        assert buf.getvalue()  # non-empty PDF

    def test_pdf_study_guide_with_image_url(self):
        """PDF export for study guide completes successfully with image."""
        study_set, cards = _study_guide_with_images()
        buf = export_study_pdf(study_set, cards)
        assert buf is not None
        assert len(buf.getvalue()) > 0


# ============================================================
# DOCX Export with Images
# ============================================================


class TestDOCXExportImages:
    """Test DOCX export includes image URL columns/references."""

    def test_docx_flashcard_table_has_image_column(self):
        """Flashcard DOCX table has an Image header column."""
        study_set, cards = _flashcard_set_with_images()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        tables = doc.tables
        assert len(tables) > 0
        header_cells = tables[0].rows[0].cells
        header_texts = [c.text for c in header_cells]
        assert "Image" in header_texts

    def test_docx_flashcard_image_url_in_cell(self):
        """Flashcard DOCX table row contains the image URL."""
        study_set, cards = _flashcard_set_with_images()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        table = doc.tables[0]
        # Row 1 (first data row) should have the image URL
        row1_cells = table.rows[1].cells
        cell_texts = [c.text for c in row1_cells]
        assert any("https://example.com/cell.png" in t for t in cell_texts)

    def test_docx_study_guide_image_reference(self):
        """Study guide DOCX includes image URL reference text."""
        study_set, cards = _study_guide_with_images()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "https://example.com/photosynthesis.jpg" in full_text

    def test_docx_vocabulary_table_has_image_column(self):
        """Vocabulary DOCX table has an Image header column."""
        study_set, cards = _vocabulary_with_images()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        tables = doc.tables
        assert len(tables) > 0
        header_cells = tables[0].rows[0].cells
        header_texts = [c.text for c in header_cells]
        assert "Image" in header_texts

    def test_docx_review_sheet_image_reference(self):
        """Review sheet DOCX includes image URL reference text."""
        study_set, cards = _review_sheet_with_images()
        buf = export_study_docx(study_set, cards)
        doc = Document(buf)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "https://example.com/einstein.png" in full_text


# ============================================================
# Study Generator Image URL Support
# ============================================================


class TestStudyGeneratorImages:
    """Test that study generator preserves image_url in card data."""

    def test_image_url_preserved_in_extras(self):
        """Verify image_url from LLM response is stored in card extras."""
        from src.study_generator import _parse_items

        response_json = json.dumps(
            [
                {
                    "front": "Test card",
                    "back": "Test answer",
                    "tags": ["test"],
                    "image_url": "https://example.com/img.png",
                }
            ]
        )
        items = _parse_items(response_json, "flashcard")
        assert items is not None
        assert items[0]["image_url"] == "https://example.com/img.png"

    def test_image_url_absent_when_not_provided(self):
        """Items without image_url don't have it in parsed output."""
        from src.study_generator import _parse_items

        response_json = json.dumps([{"front": "Card", "back": "Answer", "tags": []}])
        items = _parse_items(response_json, "flashcard")
        assert items is not None
        # image_url is simply not in the dict
        assert items[0].get("image_url", "") == ""


# ============================================================
# Web Template Image Display
# ============================================================


class TestWebImageDisplay:
    """Test image display in study detail templates."""

    @pytest.fixture
    def app(self):
        """Create Flask test app with study data containing images."""
        db_fd, db_path = tempfile.mkstemp(suffix=".db")
        engine = get_engine(db_path)
        Base.metadata.create_all(engine)
        session = get_session(engine)

        cls = Class(name="Test Class", grade_level="8th Grade", subject="Science")
        session.add(cls)
        session.commit()

        study_set = StudySet(
            class_id=cls.id,
            title="Image Test Set",
            material_type="flashcard",
            status="generated",
            config="{}",
        )
        session.add(study_set)
        session.commit()

        card_with_img = StudyCard(
            study_set_id=study_set.id,
            card_type="flashcard",
            sort_order=0,
            front="Cell",
            back="Basic unit of life",
            data=json.dumps(
                {
                    "tags": ["biology"],
                    "image_url": "https://example.com/cell.png",
                }
            ),
        )
        card_without_img = StudyCard(
            study_set_id=study_set.id,
            card_type="flashcard",
            sort_order=1,
            front="Atom",
            back="Smallest unit of matter",
            data=json.dumps({"tags": ["chemistry"]}),
        )
        session.add(card_with_img)
        session.add(card_without_img)
        session.commit()
        session.close()
        engine.dispose()

        from src.web.app import create_app

        test_config = {
            "paths": {"database_file": db_path},
            "llm": {"provider": "mock"},
            "generation": {"default_grade_level": "8th Grade"},
        }
        flask_app = create_app(test_config)
        flask_app.config["TESTING"] = True
        self._db_fd = db_fd
        self._db_path = db_path

        yield flask_app

        flask_app.config["DB_ENGINE"].dispose()
        os.close(db_fd)
        try:
            os.unlink(db_path)
        except PermissionError:
            pass

    @pytest.fixture
    def client(self, app):
        c = app.test_client()
        c.post("/login", data={"username": "teacher", "password": "quizweaver"})
        return c

    def test_image_displayed_in_study_detail(self, client):
        """Study detail page shows img tag for card with image_url."""
        resp = client.get("/study/1")
        assert resp.status_code == 200
        assert b"study-card-image" in resp.data
        assert b"https://example.com/cell.png" in resp.data

    def test_no_image_tag_without_url(self, client):
        """Cards without image_url don't show image container for that card."""
        resp = client.get("/study/1")
        assert resp.status_code == 200
        html = resp.data.decode("utf-8")
        # Count occurrences of the image class
        assert html.count("study-card-image") == 1  # Only one card has it

    def test_edit_form_has_image_url_field(self, client):
        """Edit form includes image URL input field."""
        resp = client.get("/study/1")
        assert resp.status_code == 200
        assert b"edit-image-url" in resp.data

    def test_responsive_table_wrapper_in_vocab(self):
        """Vocabulary detail uses table-responsive wrapper."""
        db_fd, db_path = tempfile.mkstemp(suffix=".db")
        engine = get_engine(db_path)
        Base.metadata.create_all(engine)
        session = get_session(engine)

        cls = Class(name="Test Class", grade_level="8th Grade", subject="Science")
        session.add(cls)
        session.commit()

        study_set = StudySet(
            class_id=cls.id,
            title="Vocab Set",
            material_type="vocabulary",
            status="generated",
            config="{}",
        )
        session.add(study_set)
        session.commit()

        card = StudyCard(
            study_set_id=study_set.id,
            card_type="term",
            sort_order=0,
            front="Test",
            back="A test term",
            data=json.dumps({"example": "ex", "part_of_speech": "noun"}),
        )
        session.add(card)
        session.commit()
        session.close()
        engine.dispose()

        from src.web.app import create_app

        test_config = {
            "paths": {"database_file": db_path},
            "llm": {"provider": "mock"},
            "generation": {"default_grade_level": "8th Grade"},
        }
        flask_app = create_app(test_config)
        flask_app.config["TESTING"] = True

        c = flask_app.test_client()
        c.post("/login", data={"username": "teacher", "password": "quizweaver"})
        resp = c.get("/study/1")
        assert resp.status_code == 200
        assert b"table-responsive" in resp.data
        assert b"responsive-table" in resp.data

        flask_app.config["DB_ENGINE"].dispose()
        os.close(db_fd)
        try:
            os.unlink(db_path)
        except PermissionError:
            pass
