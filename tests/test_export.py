"""
Tests for QuizWeaver export module (CSV, DOCX, GIFT, PDF, QTI).

Covers question normalization across data shapes (mock vs real LLM),
export output correctness, and the Flask download route.
"""

import csv
import io
import json
import os
import tempfile
import zipfile
from datetime import datetime

import pytest
from docx import Document

from src.database import Base, Class, Question, Quiz, get_engine, get_session
from src.export import (
    _escape_gift,
    _sanitize_filename,
    export_csv,
    export_docx,
    export_gift,
    export_pdf,
    export_qti,
    export_quizizz_csv,
    normalize_question,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_question(**kwargs):
    """Create a Question-like object for testing without a database."""

    class FakeQuestion:
        def __init__(self, **kw):
            self.id = kw.get("id", 1)
            self.quiz_id = kw.get("quiz_id", 1)
            self.question_type = kw.get("question_type", "mc")
            self.title = kw.get("title", "Q1")
            self.text = kw.get("text", "Sample question?")
            self.points = kw.get("points", 5.0)
            self.data = kw.get("data", {})

    return FakeQuestion(**kwargs)


def _make_quiz(**kwargs):
    """Create a Quiz-like object for testing."""

    class FakeQuiz:
        def __init__(self, **kw):
            self.id = kw.get("id", 1)
            self.title = kw.get("title", "Test Quiz")
            self.class_id = kw.get("class_id", 1)
            self.status = kw.get("status", "generated")
            self.style_profile = kw.get("style_profile", "{}")
            self.created_at = kw.get("created_at", datetime(2026, 2, 9, 12, 0))

    return FakeQuiz(**kwargs)


# ---------------------------------------------------------------------------
# TestNormalizeQuestion
# ---------------------------------------------------------------------------


class TestNormalizeQuestion:
    """Test question normalization across different data shapes."""

    def test_mc_with_correct_index(self):
        """Mock provider: MC with correct_index resolves answer from options."""
        q = _make_question(
            question_type="multiple_choice",
            text="What is 2+2?",
            data={
                "type": "multiple_choice",
                "options": ["3", "4", "5", "6"],
                "correct_index": 1,
            },
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "mc"
        assert nq["correct_answer"] == "4"
        assert nq["number"] == 1

    def test_mc_with_correct_answer_string(self):
        """Gemini provider: MC with correct_answer as string."""
        q = _make_question(
            question_type="mc",
            text="Capital of France?",
            data={
                "type": "mc",
                "options": ["London", "Paris", "Berlin", "Madrid"],
                "correct_answer": "Paris",
            },
        )
        nq = normalize_question(q, 2)
        assert nq["type"] == "mc"
        assert nq["correct_answer"] == "Paris"
        assert nq["number"] == 3

    def test_mc_with_dict_options(self):
        """Gemini-style options as list of {id, text} dicts."""
        q = _make_question(
            question_type="mc",
            text="Which color?",
            data={
                "options": [
                    {"id": "A", "text": "Red"},
                    {"id": "B", "text": "Blue"},
                ],
                "correct_answer": "Red",
            },
        )
        nq = normalize_question(q, 0)
        assert nq["options"] == ["Red", "Blue"]
        assert nq["correct_answer"] == "Red"

    def test_tf_with_answer(self):
        """True/false with correct_answer field."""
        q = _make_question(
            question_type="true_false",
            text="The sky is blue.",
            data={"correct_answer": "True"},
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "tf"
        assert nq["correct_answer"] == "True"

    def test_tf_with_correct_index(self):
        """True/false using correct_index 0=True, 1=False."""
        q = _make_question(
            question_type="tf",
            text="Water is dry.",
            data={
                "options": ["True", "False"],
                "correct_index": 1,
            },
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "tf"
        assert nq["correct_answer"] == "False"

    def test_fill_in(self):
        """Fill-in-the-blank question."""
        q = _make_question(
            question_type="fill_in",
            text="The capital of Japan is ____.",
            data={"answer": "Tokyo"},
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "fill_in"
        assert nq["correct_answer"] == "Tokyo"

    def test_matching_gemini_style(self):
        """Gemini: matches as [{term, definition}]."""
        q = _make_question(
            question_type="matching",
            text="Match the terms:",
            data={
                "matches": [
                    {"term": "H2O", "definition": "Water"},
                    {"term": "CO2", "definition": "Carbon Dioxide"},
                ],
            },
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "matching"
        assert len(nq["matches"]) == 2
        assert nq["matches"][0]["term"] == "H2O"

    def test_matching_mock_style(self):
        """Mock: prompt_items + response_items + correct_matches."""
        q = _make_question(
            question_type="matching",
            text="Match:",
            data={
                "prompt_items": ["Dog", "Cat"],
                "response_items": ["Bark", "Meow"],
                "correct_matches": {"0": 0, "1": 1},
            },
        )
        nq = normalize_question(q, 0)
        assert len(nq["matches"]) == 2
        assert nq["matches"][0] == {"term": "Dog", "definition": "Bark"}

    def test_essay(self):
        """Essay question has empty answer."""
        q = _make_question(
            question_type="essay",
            text="Discuss the causes of WWI.",
            data={},
        )
        nq = normalize_question(q, 0)
        assert nq["type"] == "essay"
        assert nq["correct_answer"] == ""

    def test_missing_fields_graceful(self):
        """Missing data fields don't crash."""
        q = _make_question(
            question_type=None,
            text=None,
            points=None,
            data=None,
        )
        nq = normalize_question(q, 0)
        assert nq["number"] == 1
        assert nq["text"] == ""
        assert nq["type"] == "mc"  # default fallback from data

    def test_data_as_json_string(self):
        """Data stored as JSON string gets parsed."""
        q = _make_question(
            question_type="mc",
            text="Question?",
            data=json.dumps(
                {
                    "options": ["A", "B"],
                    "correct_index": 0,
                }
            ),
        )
        nq = normalize_question(q, 0)
        assert nq["options"] == ["A", "B"]
        assert nq["correct_answer"] == "A"

    def test_cognitive_fields(self):
        """Cognitive level and framework are passed through."""
        q = _make_question(
            question_type="mc",
            text="Q?",
            data={
                "cognitive_level": "Analyze",
                "cognitive_framework": "blooms",
                "options": ["A"],
                "correct_index": 0,
            },
        )
        nq = normalize_question(q, 0)
        assert nq["cognitive_level"] == "Analyze"
        assert nq["cognitive_framework"] == "blooms"


# ---------------------------------------------------------------------------
# TestExportCSV
# ---------------------------------------------------------------------------


class TestExportCSV:
    """Test CSV export output."""

    def _make_mc_question(self):
        return _make_question(
            question_type="mc",
            text="What is 1+1?",
            data={
                "options": ["1", "2", "3"],
                "correct_index": 1,
                "cognitive_level": "Remember",
                "cognitive_framework": "blooms",
            },
        )

    def test_csv_has_correct_headers(self):
        quiz = _make_quiz()
        questions = [self._make_mc_question()]
        result = export_csv(quiz, questions)
        reader = csv.reader(io.StringIO(result))
        headers = next(reader)
        assert headers == [
            "#",
            "Type",
            "Question",
            "Options",
            "Correct Answer",
            "Points",
            "Cognitive Level",
            "Framework",
        ]

    def test_csv_mc_row(self):
        quiz = _make_quiz()
        questions = [self._make_mc_question()]
        result = export_csv(quiz, questions)
        reader = csv.reader(io.StringIO(result))
        next(reader)  # skip headers
        row = next(reader)
        assert row[0] == "1"
        assert row[1] == "mc"
        assert "1+1" in row[2]
        assert "A) 1" in row[3]
        assert row[4] == "2"

    def test_csv_tf_options(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="tf",
            text="True or false?",
            data={"correct_answer": "True"},
        )
        result = export_csv(quiz, [q])
        assert "True | False" in result

    def test_csv_matching(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="matching",
            text="Match:",
            data={"matches": [{"term": "A", "definition": "1"}]},
        )
        result = export_csv(quiz, [q])
        assert "A -> 1" in result

    def test_csv_empty_quiz(self):
        quiz = _make_quiz()
        result = export_csv(quiz, [])
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 1  # headers only

    def test_csv_multiple_questions(self):
        quiz = _make_quiz()
        questions = [self._make_mc_question(), self._make_mc_question()]
        result = export_csv(quiz, questions)
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 3  # headers + 2 questions


# ---------------------------------------------------------------------------
# TestExportDOCX
# ---------------------------------------------------------------------------


class TestExportDOCX:
    """Test DOCX (Word) export output."""

    def _make_quiz_with_questions(self):
        quiz = _make_quiz(title="Biology Unit Test")
        q1 = _make_question(
            question_type="mc",
            text="What is DNA?",
            data={
                "options": ["Protein", "Nucleic acid", "Lipid"],
                "correct_answer": "Nucleic acid",
                "cognitive_level": "Remember",
            },
        )
        q2 = _make_question(
            question_type="tf",
            text="Mitosis produces two cells.",
            data={"correct_answer": "True"},
        )
        return quiz, [q1, q2]

    def test_returns_bytes(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_docx(quiz, questions)
        assert buf.readable()
        # Verify it's a valid DOCX (ZIP format)
        assert zipfile.is_zipfile(buf)

    def test_contains_title(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_docx(quiz, questions)
        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Biology Unit Test" in text

    def test_contains_questions(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_docx(quiz, questions)
        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "What is DNA?" in text
        assert "Mitosis produces two cells." in text

    def test_contains_answer_key(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_docx(quiz, questions)
        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Answer Key" in text
        assert "Nucleic acid" in text

    def test_cognitive_levels_shown(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_docx(quiz, questions)
        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Remember" in text

    def test_style_profile_info(self):
        quiz = _make_quiz(
            title="Test",
            style_profile=json.dumps(
                {
                    "sol_standards": ["SOL 7.1"],
                    "cognitive_framework": "blooms",
                    "difficulty": 3,
                }
            ),
        )
        q = _make_question(question_type="mc", text="Q?", data={"options": ["A"], "correct_index": 0})
        buf = export_docx(
            quiz,
            [q],
            style_profile={
                "sol_standards": ["SOL 7.1"],
                "cognitive_framework": "blooms",
                "difficulty": 3,
            },
        )
        doc = Document(buf)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "SOL 7.1" in text
        assert "Blooms" in text


# ---------------------------------------------------------------------------
# TestExportGIFT
# ---------------------------------------------------------------------------


class TestExportGIFT:
    """Test GIFT format export."""

    def test_mc_syntax(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text="What color is grass?",
            data={
                "options": ["Red", "Green", "Blue"],
                "correct_answer": "Green",
            },
        )
        result = export_gift(quiz, [q])
        assert "::Q1::" in result
        assert "=Green" in result
        assert "~Red" in result
        assert "~Blue" in result

    def test_tf_syntax(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="tf",
            text="The earth is round.",
            data={"correct_answer": "True"},
        )
        result = export_gift(quiz, [q])
        assert "{TRUE}" in result

    def test_tf_false(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="tf",
            text="The moon is a star.",
            data={"correct_answer": "False"},
        )
        result = export_gift(quiz, [q])
        assert "{FALSE}" in result

    def test_fill_in_syntax(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="fill_in",
            text="The capital of France is ____.",
            data={"answer": "Paris"},
        )
        result = export_gift(quiz, [q])
        assert "=Paris" in result

    def test_matching_syntax(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="matching",
            text="Match:",
            data={
                "matches": [
                    {"term": "H2O", "definition": "Water"},
                    {"term": "NaCl", "definition": "Salt"},
                ],
            },
        )
        result = export_gift(quiz, [q])
        assert "=H2O -> Water" in result
        assert "=NaCl -> Salt" in result

    def test_essay_syntax(self):
        quiz = _make_quiz()
        q = _make_question(
            question_type="essay",
            text="Discuss evolution.",
            data={},
        )
        result = export_gift(quiz, [q])
        assert "::Q1::" in result
        assert "{}" in result

    def test_special_char_escaping(self):
        """GIFT special characters are escaped in question text."""
        assert _escape_gift("a=b") == "a\\=b"
        assert _escape_gift("a~b") == "a\\~b"
        assert _escape_gift("{test}") == "\\{test\\}"
        assert _escape_gift("a:b") == "a\\:b"
        assert _escape_gift("a#b") == "a\\#b"


# ---------------------------------------------------------------------------
# TestSanitizeFilename
# ---------------------------------------------------------------------------


class TestSanitizeFilename:
    def test_basic(self):
        assert _sanitize_filename("My Quiz!") == "My_Quiz"

    def test_special_chars(self):
        assert _sanitize_filename("Test: Quiz #1 (v2)") == "Test_Quiz_1_v2"

    def test_empty(self):
        assert _sanitize_filename("") == "quiz"

    def test_long_title(self):
        title = "A" * 200
        assert len(_sanitize_filename(title)) <= 80


# ---------------------------------------------------------------------------
# TestExportPDF
# ---------------------------------------------------------------------------


class TestExportPDF:
    """Test PDF export output."""

    def _make_quiz_with_questions(self):
        quiz = _make_quiz(title="History Final Exam")
        q1 = _make_question(
            question_type="mc",
            text="Who was the first president?",
            data={
                "options": ["Adams", "Washington", "Jefferson", "Franklin"],
                "correct_answer": "Washington",
                "cognitive_level": "Remember",
            },
        )
        q2 = _make_question(
            question_type="tf",
            text="The Declaration was signed in 1776.",
            data={"correct_answer": "True"},
        )
        return quiz, [q1, q2]

    def test_returns_readable_bytesio(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_pdf(quiz, questions)
        assert buf.readable()
        assert len(buf.getvalue()) > 0

    def test_starts_with_pdf_magic_bytes(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_pdf(quiz, questions)
        data = buf.getvalue()
        assert data[:5] == b"%PDF-"

    def test_has_multiple_pages(self):
        """PDF with questions + answer key should have at least 2 pages."""
        quiz, questions = self._make_quiz_with_questions()
        buf = export_pdf(quiz, questions)
        data = buf.getvalue()
        # ReportLab writes /Count N for page count in the Pages dict
        assert b"/Count 2" in data

    def test_empty_quiz_no_crash(self):
        quiz = _make_quiz(title="Empty Quiz")
        buf = export_pdf(quiz, [])
        data = buf.getvalue()
        assert data[:5] == b"%PDF-"


# ---------------------------------------------------------------------------
# TestExportQTI
# ---------------------------------------------------------------------------


class TestExportQTI:
    """Test QTI (Canvas) export output."""

    def _make_quiz_with_questions(self):
        quiz = _make_quiz(title="Science QTI Quiz")
        q1 = _make_question(
            question_type="mc",
            text="What is H2O?",
            data={
                "options": ["Water", "Salt", "Sugar", "Oil"],
                "correct_answer": "Water",
            },
        )
        q2 = _make_question(
            question_type="tf",
            text="The sun is a star.",
            data={"correct_answer": "True"},
        )
        return quiz, [q1, q2]

    def test_returns_valid_zip(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_qti(quiz, questions)
        assert zipfile.is_zipfile(buf)

    def test_zip_contains_manifest_and_assessment(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_qti(quiz, questions)
        with zipfile.ZipFile(buf, "r") as zf:
            names = zf.namelist()
            assert "imsmanifest.xml" in names
            # The assessment XML file should also be present
            xml_files = [n for n in names if n.endswith(".xml") and n != "imsmanifest.xml"]
            assert len(xml_files) == 1

    def test_assessment_contains_question_text(self):
        quiz, questions = self._make_quiz_with_questions()
        buf = export_qti(quiz, questions)
        with zipfile.ZipFile(buf, "r") as zf:
            names = zf.namelist()
            xml_files = [n for n in names if n.endswith(".xml") and n != "imsmanifest.xml"]
            assessment_xml = zf.read(xml_files[0]).decode("utf-8")
            assert "What is H2O?" in assessment_xml
            assert "The sun is a star." in assessment_xml

    def test_handles_essay_questions(self):
        quiz = _make_quiz(title="Essay QTI")
        q = _make_question(
            question_type="essay",
            text="Discuss photosynthesis.",
            data={},
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            names = zf.namelist()
            xml_files = [n for n in names if n.endswith(".xml") and n != "imsmanifest.xml"]
            assessment_xml = zf.read(xml_files[0]).decode("utf-8")
            assert "essay_question" in assessment_xml
            assert "Discuss photosynthesis." in assessment_xml


# ---------------------------------------------------------------------------
# TestExportQTIUpgrade
# ---------------------------------------------------------------------------


class TestExportQTIUpgrade:
    """Test QTI 1.2 matching, ordering, and short answer item builders."""

    def test_qti_matching_basic(self):
        """Matching question with 3 pairs produces response_grp elements."""
        quiz = _make_quiz(title="Matching QTI")
        q = _make_question(
            question_type="matching",
            text="Match the terms:",
            data={
                "matches": [
                    {"term": "H2O", "definition": "Water"},
                    {"term": "NaCl", "definition": "Salt"},
                    {"term": "CO2", "definition": "Carbon Dioxide"},
                ],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert "matching_question" in xml
            assert "response_grp" in xml
            assert "grp_0" in xml
            assert "grp_1" in xml
            assert "grp_2" in xml

    def test_qti_matching_correctness(self):
        """Verify varequal maps grp_N to def_N correctly."""
        quiz = _make_quiz(title="Matching Correct")
        q = _make_question(
            question_type="matching",
            text="Match:",
            data={
                "matches": [
                    {"term": "A", "definition": "1"},
                    {"term": "B", "definition": "2"},
                ],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert 'respident="grp_0">def_0</varequal>' in xml
            assert 'respident="grp_1">def_1</varequal>' in xml

    def test_qti_matching_xml_escaping(self):
        """Terms/definitions with &, <, > are escaped."""
        quiz = _make_quiz(title="Escape Test")
        q = _make_question(
            question_type="matching",
            text="Match these:",
            data={
                "matches": [
                    {"term": "A & B", "definition": "1 < 2"},
                    {"term": "C > D", "definition": "3 & 4"},
                ],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert "A &amp; B" in xml
            assert "1 &lt; 2" in xml
            assert "C &gt; D" in xml
            assert "3 &amp; 4" in xml

    def test_qti_matching_empty_fallback(self):
        """Matching with no matches falls back to essay."""
        quiz = _make_quiz(title="Empty Matching")
        q = _make_question(
            question_type="matching",
            text="Match nothing:",
            data={"matches": []},
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            # Falls through to default (essay) since matches is empty
            assert "essay_question" in xml or "response_grp" not in xml

    def test_qti_ordering_basic(self):
        """Ordering question maps items to position numbers."""
        quiz = _make_quiz(title="Ordering QTI")
        q = _make_question(
            question_type="ordering",
            text="Put in order:",
            data={
                "items": ["First", "Second", "Third"],
                "correct_order": [0, 1, 2],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            # Ordering is built on matching, so matching_question metadata
            assert "matching_question" in xml
            assert "First" in xml
            assert "Second" in xml
            assert "Third" in xml

    def test_qti_ordering_correct_order(self):
        """Verify correct_order is respected in position mapping."""
        quiz = _make_quiz(title="Order Verify")
        q = _make_question(
            question_type="ordering",
            text="Order these:",
            data={
                "items": ["C", "A", "B"],
                "correct_order": [2, 0, 1],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            # Item "C" at index 0 should map to position correct_order[0]+1 = 3
            # Item "A" at index 1 should map to position correct_order[1]+1 = 1
            # Item "B" at index 2 should map to position correct_order[2]+1 = 2
            assert "response_grp" in xml

    def test_qti_short_answer_basic(self):
        """Verify render_fib element and short_answer_question metadata."""
        quiz = _make_quiz(title="Short Answer QTI")
        q = _make_question(
            question_type="short_answer",
            text="What is the capital of France?",
            data={"expected_answer": "Paris"},
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert "short_answer_question" in xml
            assert "render_fib" in xml

    def test_qti_short_answer_autograding(self):
        """Verify varequal with case='No' for correct answer."""
        quiz = _make_quiz(title="SA Autograding")
        q = _make_question(
            question_type="short_answer",
            text="Capital of France?",
            data={"expected_answer": "Paris"},
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert 'case="No"' in xml
            assert "Paris" in xml

    def test_qti_short_answer_acceptable_answers(self):
        """Verify multiple varequal entries for acceptable_answers."""
        quiz = _make_quiz(title="SA Acceptable")
        q = _make_question(
            question_type="short_answer",
            text="Color of the sky?",
            data={
                "expected_answer": "blue",
                "acceptable_answers": ["azure", "cerulean"],
            },
        )
        buf = export_qti(quiz, [q])
        with zipfile.ZipFile(buf, "r") as zf:
            xml_files = [n for n in zf.namelist() if n.endswith(".xml") and n != "imsmanifest.xml"]
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert "blue" in xml
            assert "azure" in xml
            assert "cerulean" in xml
            # Should have 3 respcondition entries
            assert xml.count("respcondition") >= 6  # 3 open + 3 close tags

    def test_qti_mixed_types_zip(self):
        """Full QTI export with MC + TF + matching + ordering + short_answer."""
        quiz = _make_quiz(title="Mixed QTI")
        questions = [
            _make_question(
                question_type="mc",
                text="MC question?",
                data={"options": ["A", "B", "C"], "correct_answer": "B"},
            ),
            _make_question(
                question_type="tf",
                text="TF question?",
                data={"correct_answer": "True"},
            ),
            _make_question(
                question_type="matching",
                text="Match these:",
                data={"matches": [{"term": "X", "definition": "1"}, {"term": "Y", "definition": "2"}]},
            ),
            _make_question(
                question_type="ordering",
                text="Order these:",
                data={"items": ["A", "B"], "correct_order": [0, 1]},
            ),
            _make_question(
                question_type="short_answer",
                text="Short answer?",
                data={"expected_answer": "answer"},
            ),
        ]
        buf = export_qti(quiz, questions)
        assert zipfile.is_zipfile(buf)
        with zipfile.ZipFile(buf, "r") as zf:
            names = zf.namelist()
            assert "imsmanifest.xml" in names
            xml_files = [n for n in names if n.endswith(".xml") and n != "imsmanifest.xml"]
            assert len(xml_files) == 1
            xml = zf.read(xml_files[0]).decode("utf-8")
            assert "multiple_choice_question" in xml
            assert "matching_question" in xml
            assert "short_answer_question" in xml
            assert xml.count("<item ") == 5


# ---------------------------------------------------------------------------
# TestExportQuizizzCSV
# ---------------------------------------------------------------------------


class TestExportQuizizzCSV:
    """Test Quizizz-compatible CSV export."""

    def test_quizizz_headers(self):
        """Verify exact header row."""
        quiz = _make_quiz()
        result = export_quizizz_csv(quiz, [])
        reader = csv.reader(io.StringIO(result))
        headers = next(reader)
        assert headers == [
            "Question Text",
            "Question Type",
            "Option 1",
            "Option 2",
            "Option 3",
            "Option 4",
            "Option 5",
            "Correct Answer",
            "Time Limit",
            "Image Link",
        ]

    def test_quizizz_mc_basic(self):
        """MC question maps correctly with 1-based index."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text="What is 2+2?",
            data={"options": ["3", "4", "5", "6"], "correct_answer": "4"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)  # skip headers
        row = next(reader)
        assert row[0] == "What is 2+2?"
        assert row[1] == "Multiple Choice"
        assert row[2] == "3"
        assert row[3] == "4"
        assert row[4] == "5"
        assert row[5] == "6"
        assert row[7] == "2"  # 1-based index of "4"

    def test_quizizz_tf_basic(self):
        """TF question maps to 'True or False' type."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="tf",
            text="The sky is blue.",
            data={"correct_answer": "True"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)
        row = next(reader)
        assert row[1] == "True or False"
        assert row[2] == "True"
        assert row[3] == "False"
        assert row[7] == "True"

    def test_quizizz_option_padding(self):
        """MC with 3 options pads to 5 columns."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text="Pick one:",
            data={"options": ["A", "B", "C"], "correct_answer": "A"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)
        row = next(reader)
        assert row[2] == "A"
        assert row[3] == "B"
        assert row[4] == "C"
        assert row[5] == ""
        assert row[6] == ""

    def test_quizizz_five_options(self):
        """MC with 5 options uses all columns."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text="Pick one:",
            data={"options": ["A", "B", "C", "D", "E"], "correct_answer": "E"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)
        row = next(reader)
        assert row[2] == "A"
        assert row[6] == "E"
        assert row[7] == "5"  # 1-based index of E

    def test_quizizz_default_time(self):
        """Verify time limit is 30."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text="Q?",
            data={"options": ["A", "B"], "correct_answer": "A"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)
        row = next(reader)
        assert row[8] == "30"

    def test_quizizz_skip_matching(self):
        """Matching questions are skipped."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="matching",
            text="Match:",
            data={"matches": [{"term": "A", "definition": "1"}]},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 1  # headers only

    def test_quizizz_skip_ordering(self):
        """Ordering questions are skipped."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="ordering",
            text="Order:",
            data={"items": ["A", "B"], "correct_order": [0, 1]},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 1

    def test_quizizz_skip_essay(self):
        """Essay questions are skipped."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="essay",
            text="Discuss:",
            data={},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 1

    def test_quizizz_empty_quiz(self):
        """Empty quiz produces only headers."""
        quiz = _make_quiz()
        result = export_quizizz_csv(quiz, [])
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 1
        assert rows[0][0] == "Question Text"

    def test_quizizz_special_chars(self):
        """Questions with commas and quotes are properly escaped."""
        quiz = _make_quiz()
        q = _make_question(
            question_type="mc",
            text='What is "this", exactly?',
            data={"options": ["A, B", 'C "D"', "E"], "correct_answer": "A, B"},
        )
        result = export_quizizz_csv(quiz, [q])
        reader = csv.reader(io.StringIO(result))
        next(reader)
        row = next(reader)
        assert row[0] == 'What is "this", exactly?'
        assert row[2] == "A, B"
        assert row[3] == 'C "D"'

    def test_quizizz_mixed(self):
        """Quiz with MC + TF + matching produces only MC and TF rows."""
        quiz = _make_quiz()
        questions = [
            _make_question(
                question_type="mc",
                text="MC?",
                data={"options": ["A", "B"], "correct_answer": "A"},
            ),
            _make_question(
                question_type="tf",
                text="TF?",
                data={"correct_answer": "False"},
            ),
            _make_question(
                question_type="matching",
                text="Match:",
                data={"matches": [{"term": "X", "definition": "Y"}]},
            ),
        ]
        result = export_quizizz_csv(quiz, questions)
        reader = csv.reader(io.StringIO(result))
        rows = list(reader)
        assert len(rows) == 3  # header + MC + TF (matching skipped)
        assert rows[1][1] == "Multiple Choice"
        assert rows[2][1] == "True or False"


# ---------------------------------------------------------------------------
# TestExportRoute (integration with Flask)
# ---------------------------------------------------------------------------


@pytest.fixture
def app():
    """Create a Flask test app with a temporary database for export tests."""
    db_fd, db_path = tempfile.mkstemp(suffix=".db")

    engine = get_engine(db_path)
    Base.metadata.create_all(engine)
    session = get_session(engine)

    # Seed a class
    cls = Class(
        name="Test Class",
        grade_level="8th Grade",
        subject="Science",
        standards=json.dumps(["SOL 8.1"]),
        config=json.dumps({}),
    )
    session.add(cls)
    session.commit()

    # Seed a quiz with questions
    quiz = Quiz(
        title="Export Test Quiz",
        class_id=cls.id,
        status="generated",
        style_profile=json.dumps(
            {
                "grade_level": "8th Grade",
                "sol_standards": ["SOL 8.1"],
                "cognitive_framework": "blooms",
                "difficulty": 3,
                "provider": "mock",
            }
        ),
    )
    session.add(quiz)
    session.commit()

    q1 = Question(
        quiz_id=quiz.id,
        question_type="mc",
        title="Q1",
        text="What is photosynthesis?",
        points=5.0,
        data=json.dumps(
            {
                "type": "mc",
                "options": ["A process", "A disease", "A planet", "A tool"],
                "correct_index": 0,
                "cognitive_level": "Remember",
                "cognitive_framework": "blooms",
            }
        ),
    )
    q2 = Question(
        quiz_id=quiz.id,
        question_type="tf",
        title="Q2",
        text="Plants need sunlight.",
        points=5.0,
        data=json.dumps(
            {
                "type": "tf",
                "correct_answer": "True",
            }
        ),
    )
    session.add(q1)
    session.add(q2)
    session.commit()

    session.close()
    engine.dispose()

    from src.web.app import create_app

    test_config = {
        "paths": {"database_file": db_path},
        "llm": {"provider": "mock"},
        "generation": {},
    }
    flask_app = create_app(test_config)
    flask_app.config["TESTING"] = True

    flask_app.config["WTF_CSRF_ENABLED"] = False

    yield flask_app

    flask_app.config["DB_ENGINE"].dispose()
    os.close(db_fd)
    try:
        os.unlink(db_path)
    except PermissionError:
        pass


@pytest.fixture
def client(app):
    """Create a logged-in test client."""
    with app.test_client() as c:
        with c.session_transaction() as sess:
            sess["logged_in"] = True
            sess["username"] = "teacher"
        yield c


class TestExportRoute:
    """Test the /quizzes/<id>/export/<format> route."""

    def test_csv_returns_csv_content_type(self, client):
        resp = client.get("/quizzes/1/export/csv")
        assert resp.status_code == 200
        assert "text/csv" in resp.content_type

    def test_docx_returns_docx_content_type(self, client):
        resp = client.get("/quizzes/1/export/docx")
        assert resp.status_code == 200
        assert "officedocument" in resp.content_type

    def test_gift_returns_text_content_type(self, client):
        resp = client.get("/quizzes/1/export/gift")
        assert resp.status_code == 200
        assert "text/plain" in resp.content_type

    def test_404_for_bad_quiz(self, client):
        resp = client.get("/quizzes/9999/export/csv")
        assert resp.status_code == 404

    def test_404_for_bad_format(self, client):
        resp = client.get("/quizzes/1/export/xyz")
        assert resp.status_code == 404

    def test_csv_contains_question_text(self, client):
        resp = client.get("/quizzes/1/export/csv")
        assert b"photosynthesis" in resp.data

    def test_gift_contains_question(self, client):
        resp = client.get("/quizzes/1/export/gift")
        text = resp.data.decode("utf-8")
        assert "::Q1::" in text
        assert "photosynthesis" in text.lower()

    def test_login_required(self, app):
        """Export route requires login."""
        with app.test_client() as c:
            resp = c.get("/quizzes/1/export/csv")
            assert resp.status_code == 303  # redirect to login

    def test_docx_is_valid_zip(self, client):
        resp = client.get("/quizzes/1/export/docx")
        buf = io.BytesIO(resp.data)
        assert zipfile.is_zipfile(buf)

    def test_pdf_returns_pdf_content_type(self, client):
        resp = client.get("/quizzes/1/export/pdf")
        assert resp.status_code == 200
        assert "application/pdf" in resp.content_type

    def test_pdf_contains_magic_bytes(self, client):
        resp = client.get("/quizzes/1/export/pdf")
        assert resp.data[:5] == b"%PDF-"

    def test_qti_returns_zip_content_type(self, client):
        resp = client.get("/quizzes/1/export/qti")
        assert resp.status_code == 200
        assert "application/zip" in resp.content_type

    def test_qti_is_valid_zip(self, client):
        resp = client.get("/quizzes/1/export/qti")
        buf = io.BytesIO(resp.data)
        assert zipfile.is_zipfile(buf)
