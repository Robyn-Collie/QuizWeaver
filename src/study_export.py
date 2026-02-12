"""
Study material export module for QuizWeaver.

Exports study sets to TSV (Anki-compatible), CSV, PDF, and DOCX formats.
"""

import csv
import io
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from src.export_utils import pdf_wrap_text, sanitize_filename


def _parse_card_data(card) -> dict:
    """Parse the JSON data field from a StudyCard."""
    data = card.data
    if isinstance(data, str):
        try:
            return json.loads(data)
        except (json.JSONDecodeError, ValueError):
            return {}
    if isinstance(data, dict):
        return data
    return {}


def _sanitize_filename(title: str) -> str:
    """Sanitize a title for use as a filename.

    .. deprecated:: Use ``sanitize_filename`` from ``src.export_utils`` instead.
    """
    return sanitize_filename(title, default="study")


# ---------------------------------------------------------------------------
# TSV Export (Anki-compatible)
# ---------------------------------------------------------------------------


def export_flashcards_tsv(study_set, cards) -> str:
    """Export flashcards to Anki-compatible TSV format.

    Anki format: front<TAB>back<TAB>tags<TAB>image (no header row).
    If a card has an image_url, an <img> tag is appended for Anki compatibility.

    Args:
        study_set: StudySet ORM object.
        cards: List of StudyCard ORM objects, sorted by sort_order.

    Returns:
        TSV string (no header row).
    """
    lines = []
    for card in cards:
        front = (card.front or "").replace("\t", " ").replace("\n", " ")
        back = (card.back or "").replace("\t", " ").replace("\n", " ")
        data = _parse_card_data(card)
        tags = data.get("tags", [])
        tags_str = " ".join(str(t).replace(" ", "_") for t in tags) if tags else ""
        image_url = data.get("image_url", "")
        image_col = f'<img src="{image_url}">' if image_url else ""
        lines.append(f"{front}\t{back}\t{tags_str}\t{image_col}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# CSV Export
# ---------------------------------------------------------------------------


def export_flashcards_csv(study_set, cards) -> str:
    """Export flashcards/study material to CSV format with header row.

    Args:
        study_set: StudySet ORM object.
        cards: List of StudyCard ORM objects.

    Returns:
        CSV string with header row.
    """
    output = io.StringIO()
    writer = csv.writer(output)

    material_type = study_set.material_type

    if material_type == "flashcard":
        writer.writerow(["#", "Front", "Back", "Tags", "Image URL"])
        for i, card in enumerate(cards):
            data = _parse_card_data(card)
            tags = ", ".join(data.get("tags", []))
            writer.writerow([i + 1, card.front or "", card.back or "", tags, data.get("image_url", "")])

    elif material_type == "study_guide":
        writer.writerow(["#", "Heading", "Content", "Key Points", "Image URL"])
        for i, card in enumerate(cards):
            data = _parse_card_data(card)
            key_points = "; ".join(data.get("key_points", []))
            writer.writerow([i + 1, card.front or "", card.back or "", key_points, data.get("image_url", "")])

    elif material_type == "vocabulary":
        writer.writerow(["#", "Term", "Definition", "Example", "Part of Speech", "Image URL"])
        for i, card in enumerate(cards):
            data = _parse_card_data(card)
            writer.writerow(
                [
                    i + 1,
                    card.front or "",
                    card.back or "",
                    data.get("example", ""),
                    data.get("part_of_speech", ""),
                    data.get("image_url", ""),
                ]
            )

    elif material_type == "review_sheet":
        writer.writerow(["#", "Heading", "Content", "Type", "Image URL"])
        for i, card in enumerate(cards):
            data = _parse_card_data(card)
            writer.writerow(
                [
                    i + 1,
                    card.front or "",
                    card.back or "",
                    data.get("type", ""),
                    data.get("image_url", ""),
                ]
            )
    else:
        writer.writerow(["#", "Front", "Back", "Image URL"])
        for i, card in enumerate(cards):
            data = _parse_card_data(card)
            writer.writerow([i + 1, card.front or "", card.back or "", data.get("image_url", "")])

    return output.getvalue()


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------


def export_study_pdf(study_set, cards) -> io.BytesIO:
    """Export study material to PDF.

    Args:
        study_set: StudySet ORM object.
        cards: List of StudyCard ORM objects.

    Returns:
        BytesIO buffer containing the PDF file.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    # Title
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2, y, study_set.title or "Study Material")
    y -= 24

    # Type badge
    c.setFont("Helvetica", 10)
    type_label = (study_set.material_type or "").replace("_", " ").title()
    c.drawCentredString(width / 2, y, f"Type: {type_label}")
    y -= 30

    material_type = study_set.material_type

    for i, card in enumerate(cards):
        data = _parse_card_data(card)

        # Check for page break
        if y < 120:
            c.showPage()
            y = height - 50

        if material_type == "flashcard":
            _pdf_draw_flashcard(c, card, data, i, y, width, height)
            y -= 60
        elif material_type == "study_guide":
            y = _pdf_draw_study_guide_section(c, card, data, y, width, height)
        elif material_type == "vocabulary":
            y = _pdf_draw_vocab_entry(c, card, data, i, y, width, height)
        elif material_type == "review_sheet":
            y = _pdf_draw_review_item(c, card, data, y, width, height)
        else:
            _pdf_draw_flashcard(c, card, data, i, y, width, height)
            y -= 60

    c.save()
    buf.seek(0)
    return buf


def _pdf_draw_flashcard(c, card, data, index, y, width, height):
    """Draw a flashcard on the PDF."""
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, f"{index + 1}. {card.front or ''}")
    c.setFont("Helvetica", 10)
    y_back = y - 16
    c.drawString(70, y_back, card.back or "")
    tags = data.get("tags", [])
    if tags:
        c.setFont("Helvetica-Oblique", 8)
        c.drawString(70, y_back - 14, f"Tags: {', '.join(str(t) for t in tags)}")
    image_url = data.get("image_url", "")
    if image_url:
        c.setFont("Helvetica-Oblique", 7)
        tag_offset = 14 if tags else 0
        c.drawString(70, y_back - 14 - tag_offset, f"Image: {image_url}")


def _pdf_draw_study_guide_section(c, card, data, y, width, height):
    """Draw a study guide section. Returns new y position."""
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, card.front or "Section")
    y -= 18

    c.setFont("Helvetica", 10)
    y = _pdf_wrap_text(c, card.back or "", 50, y, width - 100, height)
    y -= 6

    key_points = data.get("key_points", [])
    if key_points:
        c.setFont("Helvetica-Bold", 9)
        c.drawString(60, y, "Key Points:")
        y -= 14
        c.setFont("Helvetica", 9)
        for point in key_points:
            if y < 60:
                c.showPage()
                y = height - 50
            c.drawString(70, y, f"- {point}")
            y -= 13

    y -= 12
    return y


def _pdf_draw_vocab_entry(c, card, data, index, y, width, height):
    """Draw a vocabulary entry. Returns new y position."""
    c.setFont("Helvetica-Bold", 11)
    pos = data.get("part_of_speech", "")
    pos_str = f" ({pos})" if pos else ""
    c.drawString(50, y, f"{index + 1}. {card.front or ''}{pos_str}")
    y -= 16

    c.setFont("Helvetica", 10)
    c.drawString(70, y, card.back or "")
    y -= 14

    example = data.get("example", "")
    if example:
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(70, y, f"Example: {example}")
        y -= 14

    y -= 8
    return y


def _pdf_draw_review_item(c, card, data, y, width, height):
    """Draw a review sheet item. Returns new y position."""
    item_type = data.get("type", "")
    type_str = f" [{item_type.upper()}]" if item_type else ""

    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, f"{card.front or ''}{type_str}")
    y -= 16

    c.setFont("Helvetica", 10)
    y = _pdf_wrap_text(c, card.back or "", 60, y, width - 120, height)
    y -= 12
    return y


_pdf_wrap_text = pdf_wrap_text  # Backward-compatible alias


# ---------------------------------------------------------------------------
# DOCX (Word) Export
# ---------------------------------------------------------------------------


def export_study_docx(study_set, cards) -> io.BytesIO:
    """Export study material to a Word document.

    Args:
        study_set: StudySet ORM object.
        cards: List of StudyCard ORM objects.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Title
    title_p = doc.add_heading(study_set.title or "Study Material", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Type info
    type_label = (study_set.material_type or "").replace("_", " ").title()
    p = doc.add_paragraph(f"Type: {type_label}")
    p.style.font.size = Pt(10)

    material_type = study_set.material_type

    if material_type == "flashcard":
        _docx_flashcards(doc, cards)
    elif material_type == "study_guide":
        _docx_study_guide(doc, cards)
    elif material_type == "vocabulary":
        _docx_vocabulary(doc, cards)
    elif material_type == "review_sheet":
        _docx_review_sheet(doc, cards)
    else:
        _docx_flashcards(doc, cards)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_flashcards(doc, cards):
    """Add flashcards as a table to the Word document."""
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Front"
    hdr[2].text = "Back"
    hdr[3].text = "Image"

    for i, card in enumerate(cards):
        data = _parse_card_data(card)
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = card.front or ""
        row[2].text = card.back or ""
        image_url = data.get("image_url", "")
        row[3].text = image_url if image_url else ""


def _docx_study_guide(doc, cards):
    """Add study guide sections to the Word document."""
    for card in cards:
        data = _parse_card_data(card)
        doc.add_heading(card.front or "Section", level=2)
        doc.add_paragraph(card.back or "")

        image_url = data.get("image_url", "")
        if image_url:
            p = doc.add_paragraph()
            run = p.add_run(f"Image: {image_url}")
            run.italic = True
            run.font.size = Pt(8)

        key_points = data.get("key_points", [])
        if key_points:
            p = doc.add_paragraph()
            run = p.add_run("Key Points:")
            run.bold = True
            run.font.size = Pt(10)
            for point in key_points:
                doc.add_paragraph(point, style="List Bullet")


def _docx_vocabulary(doc, cards):
    """Add vocabulary as a table to the Word document."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Term"
    hdr[1].text = "Definition"
    hdr[2].text = "Example"
    hdr[3].text = "Part of Speech"
    hdr[4].text = "Image"

    for card in cards:
        data = _parse_card_data(card)
        row = table.add_row().cells
        row[0].text = card.front or ""
        row[1].text = card.back or ""
        row[2].text = data.get("example", "")
        row[3].text = data.get("part_of_speech", "")
        image_url = data.get("image_url", "")
        row[4].text = image_url if image_url else ""


def _docx_review_sheet(doc, cards):
    """Add review sheet items to the Word document."""
    for card in cards:
        data = _parse_card_data(card)
        item_type = data.get("type", "")
        type_str = f" [{item_type.upper()}]" if item_type else ""

        p = doc.add_paragraph()
        run = p.add_run(f"{card.front or ''}{type_str}")
        run.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph(card.back or "")

        image_url = data.get("image_url", "")
        if image_url:
            p = doc.add_paragraph()
            run = p.add_run(f"Image: {image_url}")
            run.italic = True
            run.font.size = Pt(8)
