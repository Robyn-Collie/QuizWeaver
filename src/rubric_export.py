"""
Rubric export module for QuizWeaver.

Exports rubrics to PDF, DOCX (Word), and CSV formats.
"""

import csv
import io
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from src.export_utils import pdf_wrap_text

PROFICIENCY_LABELS = ["Beginning", "Developing", "Proficient", "Advanced"]


def _parse_levels(criterion) -> list:
    """Parse the levels JSON field from a RubricCriterion."""
    levels = criterion.levels
    if isinstance(levels, str):
        try:
            return json.loads(levels)
        except (json.JSONDecodeError, ValueError):
            return []
    if isinstance(levels, list):
        return levels
    return []


# ---------------------------------------------------------------------------
# CSV Export
# ---------------------------------------------------------------------------


def export_rubric_csv(rubric, criteria) -> str:
    """Export rubric to CSV format.

    Args:
        rubric: Rubric ORM object.
        criteria: List of RubricCriterion ORM objects, sorted by sort_order.

    Returns:
        CSV string with header row and criteria rows.
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # Header
    header = ["Criterion", "Description", "Max Points"]
    for label in PROFICIENCY_LABELS:
        header.append(label)
    writer.writerow(header)

    for c in criteria:
        levels = _parse_levels(c)
        level_descs = {}
        for lv in levels:
            label = lv.get("label", "")
            level_descs[label] = lv.get("description", "")

        row = [
            c.criterion or "",
            c.description or "",
            c.max_points or 0,
        ]
        for label in PROFICIENCY_LABELS:
            row.append(level_descs.get(label, ""))
        writer.writerow(row)

    return output.getvalue()


# ---------------------------------------------------------------------------
# DOCX (Word) Export
# ---------------------------------------------------------------------------


def export_rubric_docx(rubric, criteria) -> io.BytesIO:
    """Export rubric to a Word document with a formatted table.

    Args:
        rubric: Rubric ORM object.
        criteria: List of RubricCriterion ORM objects.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Title
    title_p = doc.add_heading(rubric.title or "Rubric", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Table: Criterion | Max Pts | Beginning | Developing | Proficient | Advanced
    col_count = 2 + len(PROFICIENCY_LABELS)  # criterion + max_points + 4 levels
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Criterion"
    hdr[1].text = "Max Pts"
    for i, label in enumerate(PROFICIENCY_LABELS):
        hdr[2 + i].text = label

    # Make header bold
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for c in criteria:
        levels = _parse_levels(c)
        level_descs = {}
        for lv in levels:
            label = lv.get("label", "")
            level_descs[label] = lv.get("description", "")

        row_cells = table.add_row().cells
        # Criterion name + description
        p = row_cells[0].paragraphs[0]
        run = p.add_run(c.criterion or "")
        run.bold = True
        run.font.size = Pt(9)
        if c.description:
            p.add_run(f"\n{c.description}").font.size = Pt(8)

        row_cells[1].text = str(c.max_points or 0)
        for para in row_cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, label in enumerate(PROFICIENCY_LABELS):
            row_cells[2 + i].text = level_descs.get(label, "")
            for para in row_cells[2 + i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # Total points row
    total_pts = sum(c.max_points or 0 for c in criteria)
    p = doc.add_paragraph()
    run = p.add_run(f"Total Points: {total_pts}")
    run.bold = True
    run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------


def export_rubric_pdf(rubric, criteria) -> io.BytesIO:
    """Export rubric to a PDF with a table layout.

    Args:
        rubric: Rubric ORM object.
        criteria: List of RubricCriterion ORM objects.

    Returns:
        BytesIO buffer containing the PDF file.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, y, rubric.title or "Rubric")
    y -= 30

    # Stacked vertical layout — one block per criterion
    left_margin = 50
    max_text_width = width - 100  # 50px margins each side

    for cr in criteria:
        levels = _parse_levels(cr)
        level_descs = {}
        for lv in levels:
            label = lv.get("label", "")
            level_descs[label] = lv.get("description", "")

        # Page break check (rough estimate: header + description + 4 levels)
        if y < 160:
            c.showPage()
            y = height - 50

        # Criterion name + points
        c.setFont("Helvetica-Bold", 10)
        pts_str = f" ({cr.max_points or 0} pts)" if cr.max_points else ""
        y = pdf_wrap_text(c, f"{cr.criterion or 'Criterion'}{pts_str}", left_margin, y, max_text_width, height)

        # Description
        if cr.description:
            c.setFont("Helvetica-Oblique", 8)
            y = pdf_wrap_text(c, cr.description, left_margin + 10, y, max_text_width - 10, height)
            y -= 4

        # Proficiency levels
        for label in PROFICIENCY_LABELS:
            desc = level_descs.get(label, "")
            if not desc:
                continue
            if y < 60:
                c.showPage()
                y = height - 50
            c.setFont("Helvetica-Bold", 8)
            c.drawString(left_margin + 10, y, f"{label}:")
            c.setFont("Helvetica", 8)
            y = pdf_wrap_text(c, desc, left_margin + 80, y, max_text_width - 80, height)

        y -= 8
        # Separator line
        c.setStrokeColorRGB(0.85, 0.85, 0.85)
        c.line(left_margin, y + 4, width - 50, y + 4)
        c.setStrokeColorRGB(0, 0, 0)
        y -= 8

    # Total
    if y < 60:
        c.showPage()
        y = height - 50
    total_pts = sum(cr.max_points or 0 for cr in criteria)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(left_margin, y, f"Total Points: {total_pts}")

    c.save()
    buf.seek(0)
    return buf
