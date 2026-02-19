"""
Pacing guide export module for QuizWeaver.

Exports pacing guides to CSV, PDF, and DOCX (Word) formats.
"""

import csv
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from src.export_utils import parse_json_field, sanitize_csv_cell

# ---------------------------------------------------------------------------
# CSV Export
# ---------------------------------------------------------------------------


def export_pacing_csv(guide, units) -> str:
    """Export pacing guide to CSV.

    Columns: Unit, Title, Weeks, Standards, Topics, Assessment, Notes.

    Args:
        guide: PacingGuide ORM object.
        units: List of PacingGuideUnit ORM objects.

    Returns:
        CSV string.
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow(
        ["Unit", "Title", "Weeks", "Standards", "Topics", "Assessment", "Notes"]
    )

    for unit in units:
        standards = parse_json_field(unit.standards, fallback=[])
        topics = parse_json_field(unit.topics, fallback=[])

        writer.writerow(
            [
                sanitize_csv_cell(str(unit.unit_number)),
                sanitize_csv_cell(unit.title or ""),
                sanitize_csv_cell(f"{unit.start_week}-{unit.end_week}"),
                sanitize_csv_cell(", ".join(standards) if standards else ""),
                sanitize_csv_cell(", ".join(topics) if topics else ""),
                sanitize_csv_cell(unit.assessment_type or ""),
                sanitize_csv_cell(unit.notes or ""),
            ]
        )

    return output.getvalue()


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------


def export_pacing_pdf(guide, units) -> io.BytesIO:
    """Export pacing guide to PDF with a table layout.

    Args:
        guide: PacingGuide ORM object.
        units: List of PacingGuideUnit ORM objects.

    Returns:
        BytesIO buffer containing the PDF file.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    # Title
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2, y, guide.title or "Pacing Guide")
    y -= 24

    # Metadata
    c.setFont("Helvetica", 10)
    meta_parts = []
    if guide.school_year:
        meta_parts.append(f"School Year: {guide.school_year}")
    meta_parts.append(f"Total Weeks: {guide.total_weeks}")
    if meta_parts:
        c.drawCentredString(width / 2, y, " | ".join(meta_parts))
        y -= 20

    # Table header
    if y < 100:
        c.showPage()
        y = height - 50

    c.setFont("Helvetica-Bold", 9)
    col_x = [40, 80, 200, 290, 380, 470, 530]
    headers = ["Unit", "Title", "Weeks", "Standards", "Topics", "Type", "Notes"]
    for i, header in enumerate(headers):
        c.drawString(col_x[i], y, header)
    y -= 4
    c.setStrokeColor(colors.grey)
    c.line(40, y, width - 40, y)
    y -= 14

    # Table rows
    c.setFont("Helvetica", 8)
    for unit in units:
        if y < 60:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica-Bold", 9)
            for i, header in enumerate(headers):
                c.drawString(col_x[i], y, header)
            y -= 4
            c.line(40, y, width - 40, y)
            y -= 14
            c.setFont("Helvetica", 8)

        standards = parse_json_field(unit.standards, fallback=[])
        topics = parse_json_field(unit.topics, fallback=[])

        row = [
            str(unit.unit_number),
            (unit.title or "")[:18],
            f"Wk {unit.start_week}-{unit.end_week}",
            ", ".join(standards)[:20] if standards else "",
            ", ".join(topics)[:20] if topics else "",
            (unit.assessment_type or "")[:10],
            (unit.notes or "")[:20],
        ]
        for i, cell in enumerate(row):
            c.drawString(col_x[i], y, cell)
        y -= 14

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# DOCX (Word) Export
# ---------------------------------------------------------------------------


def export_pacing_docx(guide, units) -> io.BytesIO:
    """Export pacing guide to a Word document with a formatted table.

    Args:
        guide: PacingGuide ORM object.
        units: List of PacingGuideUnit ORM objects.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Title
    title_p = doc.add_heading(guide.title or "Pacing Guide", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_parts = []
    if guide.school_year:
        meta_parts.append(f"School Year: {guide.school_year}")
    meta_parts.append(f"Total Weeks: {guide.total_weeks}")
    if meta_parts:
        p = doc.add_paragraph(" | ".join(meta_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    # Table
    headers = ["Unit", "Title", "Weeks", "Standards", "Topics", "Assessment", "Notes"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for unit in units:
        standards = parse_json_field(unit.standards, fallback=[])
        topics = parse_json_field(unit.topics, fallback=[])

        row_cells = table.add_row().cells
        row_cells[0].text = str(unit.unit_number)
        row_cells[1].text = unit.title or ""
        row_cells[2].text = f"Wk {unit.start_week}-{unit.end_week}"
        row_cells[3].text = ", ".join(standards) if standards else ""
        row_cells[4].text = ", ".join(topics) if topics else ""
        row_cells[5].text = unit.assessment_type or ""
        row_cells[6].text = unit.notes or ""

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
