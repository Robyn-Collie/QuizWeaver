"""
Lesson plan export module for QuizWeaver.

Exports lesson plans to PDF and DOCX (Word) formats.
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from src.export_utils import parse_json_field, pdf_wrap_text, sanitize_filename

SECTION_LABELS = {
    "learning_objectives": "Learning Objectives",
    "materials_needed": "Materials Needed",
    "warm_up": "Warm-Up (5-10 min)",
    "direct_instruction": "Direct Instruction (10-15 min)",
    "guided_practice": "Guided Practice (10-15 min)",
    "independent_practice": "Independent Practice (10-15 min)",
    "assessment": "Assessment / Check for Understanding (5 min)",
    "closure": "Closure (3-5 min)",
    "differentiation": "Differentiation",
    "standards_alignment": "Standards Alignment",
}

SECTION_ORDER = [
    "learning_objectives",
    "materials_needed",
    "warm_up",
    "direct_instruction",
    "guided_practice",
    "independent_practice",
    "assessment",
    "closure",
    "differentiation",
    "standards_alignment",
]


def _parse_plan_data(lesson_plan) -> dict:
    """Parse the JSON plan_data field from a LessonPlan."""
    return parse_json_field(lesson_plan.plan_data)


def _sanitize_filename(title: str) -> str:
    """Sanitize a title for use as a filename.

    .. deprecated:: Use ``sanitize_filename`` from ``src.export_utils`` instead.
    """
    return sanitize_filename(title, default="lesson_plan")


# Backward-compatible alias for internal call sites
_pdf_wrap_text = pdf_wrap_text


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------


def export_lesson_plan_pdf(lesson_plan) -> io.BytesIO:
    """Export a lesson plan to PDF.

    Args:
        lesson_plan: LessonPlan ORM object.

    Returns:
        BytesIO buffer containing the PDF file.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    # Title
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(width / 2, y, lesson_plan.title or "Lesson Plan")
    y -= 24

    # Metadata line
    c.setFont("Helvetica", 10)
    meta_parts = []
    if lesson_plan.grade_level:
        meta_parts.append(f"Grade: {lesson_plan.grade_level}")
    if lesson_plan.duration_minutes:
        meta_parts.append(f"Duration: {lesson_plan.duration_minutes} min")
    if meta_parts:
        c.drawCentredString(width / 2, y, " | ".join(meta_parts))
        y -= 16

    # AI draft notice
    c.setFont("Helvetica-Oblique", 8)
    c.drawCentredString(width / 2, y, "AI-Generated Draft - Review and edit all sections before classroom use.")
    y -= 24

    # Sections
    plan_data = _parse_plan_data(lesson_plan)

    for section_key in SECTION_ORDER:
        content = plan_data.get(section_key, "")
        if not content:
            continue

        label = SECTION_LABELS.get(section_key, section_key.replace("_", " ").title())

        # Check page break
        if y < 100:
            c.showPage()
            y = height - 50

        # Section heading
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, label)
        y -= 18

        # Section content
        c.setFont("Helvetica", 10)
        y = _pdf_wrap_text(c, content, 60, y, width - 120, height)
        y -= 12

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# DOCX (Word) Export
# ---------------------------------------------------------------------------


def export_lesson_plan_docx(lesson_plan) -> io.BytesIO:
    """Export a lesson plan to a Word document.

    Args:
        lesson_plan: LessonPlan ORM object.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    doc = Document()

    # Title
    title_p = doc.add_heading(lesson_plan.title or "Lesson Plan", level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_parts = []
    if lesson_plan.grade_level:
        meta_parts.append(f"Grade Level: {lesson_plan.grade_level}")
    if lesson_plan.duration_minutes:
        meta_parts.append(f"Duration: {lesson_plan.duration_minutes} minutes")
    if meta_parts:
        p = doc.add_paragraph(" | ".join(meta_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    # AI draft notice
    p = doc.add_paragraph("AI-Generated Draft - Review and edit all sections before classroom use.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.italic = True

    # Sections
    plan_data = _parse_plan_data(lesson_plan)

    for section_key in SECTION_ORDER:
        content = plan_data.get(section_key, "")
        if not content:
            continue

        label = SECTION_LABELS.get(section_key, section_key.replace("_", " ").title())

        doc.add_heading(label, level=2)
        doc.add_paragraph(content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
