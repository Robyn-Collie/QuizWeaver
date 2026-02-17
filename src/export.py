"""
Quiz export module for QuizWeaver.

Exports quizzes to CSV, DOCX (Word), GIFT (Moodle), PDF, and QTI (Canvas) formats.
Handles normalization of different question data shapes from
mock vs real LLM providers.
"""

import csv
import io
import json
import os
import uuid
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from src.export_utils import sanitize_csv_cell, sanitize_filename

# Type normalization map: long form -> short form
TYPE_MAP = {
    "multiple_choice": "mc",
    "true_false": "tf",
    "short_answer": "short_answer",
    "fill_in": "fill_in",
    "fill_in_blank": "fill_in",
    "fill_in_the_blank": "fill_in",
    "matching": "matching",
    "ordering": "ordering",
    "essay": "essay",
}


def normalize_question(question_obj, index: int) -> Dict[str, Any]:
    """Normalize a Question ORM object into a clean dict.

    Handles different data shapes from mock vs real LLM providers:
    - Type field: "multiple_choice" vs "mc", "true_false" vs "tf"
    - Answer: correct_index (int) vs correct_answer/answer (string)
    - Matching: prompt_items/response_items vs matches [{term, definition}]
    - Text: "text" vs "question" field

    Args:
        question_obj: Question ORM object with .data, .text, .question_type, etc.
        index: 0-based question index for numbering.

    Returns:
        Normalized dict with consistent field names.
    """
    # Parse data field if stored as JSON string
    data = question_obj.data
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except (json.JSONDecodeError, ValueError):
            data = {}
    if not isinstance(data, dict):
        data = {}

    # Resolve type
    raw_type = question_obj.question_type or data.get("type", "mc")
    q_type = TYPE_MAP.get(raw_type, raw_type)

    # Resolve text: prefer ORM .text, then data.text, then data.question
    text = question_obj.text or data.get("text") or data.get("question", "")

    # Resolve options
    options = data.get("options", [])
    # Handle Gemini-style options as list of dicts {id, text}
    if options and isinstance(options[0], dict):
        options = [opt.get("text", str(opt)) for opt in options]

    # Resolve correct answer
    correct_answer = _resolve_correct_answer(data, options, q_type)

    # Resolve matching pairs
    matches = _resolve_matches(data)

    # Resolve ordering fields
    ordering_items = data.get("items", [])
    ordering_correct_order = data.get("correct_order", [])
    ordering_instructions = data.get("instructions", "")

    # Resolve short answer fields
    expected_answer = data.get("expected_answer", "")
    acceptable_answers = data.get("acceptable_answers", [])
    rubric_hint = data.get("rubric_hint", "")

    # For short_answer, use expected_answer as correct_answer if not already set
    if q_type == "short_answer" and not correct_answer and expected_answer:
        correct_answer = expected_answer

    return {
        "number": index + 1,
        "type": q_type,
        "text": text,
        "options": options,
        "correct_answer": correct_answer,
        "matches": matches,
        "ordering_items": ordering_items,
        "ordering_correct_order": ordering_correct_order,
        "ordering_instructions": ordering_instructions,
        "expected_answer": expected_answer,
        "acceptable_answers": acceptable_answers,
        "rubric_hint": rubric_hint,
        "word_bank": data.get("word_bank"),
        "points": question_obj.points or data.get("points", 0),
        "cognitive_level": data.get("cognitive_level"),
        "cognitive_framework": data.get("cognitive_framework"),
        "image_description": data.get("image_description") or data.get("image_ref"),
    }


def _resolve_correct_answer(data: dict, options: list, q_type: str) -> str:
    """Resolve the correct answer string from various data shapes."""
    # Try direct answer fields first
    answer = data.get("correct_answer") or data.get("answer")
    if answer is not None:
        return str(answer)

    # Fall back to correct_index into options
    correct_index = data.get("correct_index")
    if correct_index is not None and options:
        try:
            idx = int(correct_index)
            if 0 <= idx < len(options):
                return str(options[idx])
        except (ValueError, TypeError):
            pass

    # For T/F questions, try is_true field
    if q_type == "tf":
        is_true = data.get("is_true")
        if is_true is not None:
            return "True" if is_true else "False"

    return ""


def _resolve_matches(data: dict) -> List[Dict[str, str]]:
    """Resolve matching pairs from various data shapes."""
    # Gemini style: matches = [{term, definition}]
    matches = data.get("matches")
    if matches and isinstance(matches, list):
        result = []
        for m in matches:
            if isinstance(m, dict):
                result.append(
                    {
                        "term": m.get("term", ""),
                        "definition": m.get("definition", ""),
                    }
                )
        if result:
            return result

    # Mock style: prompt_items + response_items + correct_matches
    prompts = data.get("prompt_items", [])
    responses = data.get("response_items", [])
    correct_matches = data.get("correct_matches", {})
    if prompts and responses:
        result = []
        for i, prompt in enumerate(prompts):
            match_idx = correct_matches.get(str(i), i)
            definition = responses[match_idx] if match_idx < len(responses) else ""
            result.append({"term": prompt, "definition": definition})
        return result

    return []


def _escape_gift(text: str) -> str:
    """Escape special GIFT format characters."""
    for char in ["~", "=", "#", "{", "}", ":"]:
        text = text.replace(char, "\\" + char)
    return text


def _sanitize_filename(title: str) -> str:
    """Sanitize a quiz title for use as a filename.

    .. deprecated:: Use ``sanitize_filename`` from ``src.export_utils`` instead.
    """
    return sanitize_filename(title, default="quiz")


# ---------------------------------------------------------------------------
# CSV Export
# ---------------------------------------------------------------------------


def export_csv(
    quiz, questions, style_profile: Optional[dict] = None, student_mode: bool = False
) -> str:
    """Export quiz questions to CSV format.

    Args:
        quiz: Quiz ORM object.
        questions: List of Question ORM objects.
        style_profile: Parsed style profile dict (optional).
        student_mode: If True, omit correct answer and cognitive columns.

    Returns:
        CSV string with headers and question rows.
    """
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    if student_mode:
        writer.writerow(["#", "Type", "Question", "Options", "Points"])
    else:
        writer.writerow(
            [
                "#",
                "Type",
                "Question",
                "Options",
                "Correct Answer",
                "Points",
                "Cognitive Level",
                "Framework",
            ]
        )

    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        options_str = _format_options_csv(nq)
        if student_mode:
            writer.writerow(
                [
                    nq["number"],
                    nq["type"],
                    sanitize_csv_cell(nq["text"]),
                    sanitize_csv_cell(options_str),
                    nq["points"],
                ]
            )
        else:
            writer.writerow(
                [
                    nq["number"],
                    nq["type"],
                    sanitize_csv_cell(nq["text"]),
                    sanitize_csv_cell(options_str),
                    sanitize_csv_cell(nq["correct_answer"]),
                    nq["points"],
                    nq["cognitive_level"] or "",
                    nq["cognitive_framework"] or "",
                ]
            )

    return output.getvalue()


def _format_options_csv(nq: dict) -> str:
    """Format options for CSV column."""
    if nq["type"] == "matching" and nq["matches"]:
        return " | ".join(f"{m['term']} -> {m['definition']}" for m in nq["matches"])
    if nq["type"] == "ordering" and nq.get("ordering_items"):
        return ", ".join(nq["ordering_items"])
    if nq["type"] == "short_answer":
        answers = nq.get("acceptable_answers", [])
        if answers:
            return " | ".join(answers)
        return nq.get("expected_answer", "")
    if nq["type"] == "tf":
        return "True | False"
    if nq["options"]:
        letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        parts = []
        for j, opt in enumerate(nq["options"]):
            letter = letters[j] if j < len(letters) else str(j)
            parts.append(f"{letter}) {opt}")
        return " | ".join(parts)
    return ""


# ---------------------------------------------------------------------------
# Quizizz-Compatible CSV Export
# ---------------------------------------------------------------------------


def export_quizizz_csv(quiz, questions, style_profile: Optional[dict] = None) -> str:
    """Export quiz in Quizizz-compatible CSV import format.

    Quizizz import format columns:
    Question Text, Question Type, Option 1, Option 2, Option 3, Option 4, Option 5,
    Correct Answer, Time Limit, Image Link

    Only MC and TF questions are supported by Quizizz. Other types are skipped.
    """
    import logging

    logger = logging.getLogger(__name__)

    output = io.StringIO()
    writer = csv.writer(output)

    # Header row (Quizizz official import format)
    writer.writerow([
        "Question Text", "Question Type",
        "Option 1", "Option 2", "Option 3", "Option 4", "Option 5",
        "Correct Answer", "Time Limit", "Image Link",
    ])

    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        q_type = nq["type"]

        if q_type == "mc":
            options = nq["options"][:5]  # Quizizz supports up to 5 options
            # Pad to 5 options
            while len(options) < 5:
                options.append("")

            # Find correct answer as 1-based index
            correct_answer = ""
            for idx, opt in enumerate(nq["options"][:5]):
                if str(opt) == nq["correct_answer"]:
                    correct_answer = str(idx + 1)
                    break
            if not correct_answer and nq["correct_answer"]:
                correct_answer = nq["correct_answer"]

            writer.writerow([
                sanitize_csv_cell(nq["text"]), "Multiple Choice",
                sanitize_csv_cell(options[0]), sanitize_csv_cell(options[1]),
                sanitize_csv_cell(options[2]), sanitize_csv_cell(options[3]),
                sanitize_csv_cell(options[4]),
                sanitize_csv_cell(correct_answer), 30, "",
            ])

        elif q_type == "tf":
            writer.writerow([
                sanitize_csv_cell(nq["text"]), "True or False",
                "True", "False", "", "", "",
                nq["correct_answer"] or "True", 30, "",
            ])

        else:
            # matching, ordering, short_answer, fill_in, essay -- skip with warning
            logger.debug(
                "Quizizz CSV: skipping question %d (type=%s, not supported by Quizizz)",
                nq["number"], q_type,
            )

    return output.getvalue()


# ---------------------------------------------------------------------------
# DOCX (Word) Export
# ---------------------------------------------------------------------------


def export_docx(
    quiz, questions, style_profile: Optional[dict] = None, student_mode: bool = False
) -> io.BytesIO:
    """Export quiz to a Word document.

    Args:
        quiz: Quiz ORM object.
        questions: List of Question ORM objects.
        style_profile: Parsed style profile dict (optional).
        student_mode: If True, suppress correct-answer highlighting,
            cognitive levels, image descriptions, and move answer key
            to a separate page without inline hints.

    Returns:
        BytesIO buffer containing the .docx file.
    """
    if style_profile is None:
        style_profile = {}

    doc = Document()

    # Title
    title = doc.add_heading(quiz.title or "Quiz", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info block (skip provider/generated-by in student mode)
    _add_docx_info(doc, quiz, style_profile, student_mode=student_mode)

    # Name / Date line for student copies
    if student_mode:
        doc.add_paragraph("Name: ____________________________  Date: ____________")

    # Questions section
    doc.add_heading("Questions", level=2)

    normalized = []
    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        normalized.append(nq)
        _add_docx_question(doc, nq, student_mode=student_mode)

    # Answer key (new page) — only in teacher mode
    if not student_mode:
        doc.add_page_break()
        doc.add_heading("Answer Key", level=2)
        _add_docx_answer_key(doc, normalized)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_docx_info(doc, quiz, style_profile: dict, student_mode: bool = False):
    """Add quiz metadata info block to the document."""
    info_lines = []

    sol = style_profile.get("sol_standards")
    if sol:
        if isinstance(sol, list):
            sol = ", ".join(sol)
        info_lines.append(f"Standards: {sol}")

    if not student_mode:
        framework = style_profile.get("cognitive_framework")
        if framework:
            info_lines.append(f"Framework: {framework.capitalize()}")

        difficulty = style_profile.get("difficulty")
        if difficulty:
            info_lines.append(f"Difficulty: {difficulty}/5")

        provider = style_profile.get("provider")
        model = style_profile.get("model")
        if provider:
            generated_by = model if model else provider
            info_lines.append(f"Generated by: {generated_by}")

    created = getattr(quiz, "created_at", None)
    if created:
        if isinstance(created, datetime):
            info_lines.append(f"Date: {created.strftime('%Y-%m-%d')}")
        else:
            info_lines.append(f"Date: {created}")

    if info_lines:
        for line in info_lines:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(10)


def _add_docx_question(doc, nq: dict, student_mode: bool = False):
    """Add a single question to the Word document."""
    # Question header: "1. [MC] (5 pts) - Remember"
    header_parts = [f"{nq['number']}."]
    header_parts.append(f"[{nq['type'].upper()}]")
    header_parts.append(f"({nq['points']} pts)")
    if nq["cognitive_level"] and not student_mode:
        header_parts.append(f"- {nq['cognitive_level']}")

    p = doc.add_paragraph()
    run = p.add_run(" ".join(header_parts))
    run.bold = True
    run.font.size = Pt(11)

    # Question text
    doc.add_paragraph(nq["text"])

    # Image description — teacher mode only
    if not student_mode and nq.get("image_description"):
        p = doc.add_paragraph()
        run = p.add_run(f"Suggested image: {nq['image_description']}")
        run.italic = True
        run.font.size = Pt(9)

    # Answer area based on type
    if nq["type"] == "mc":
        _add_docx_mc_options(doc, nq, student_mode=student_mode)
    elif nq["type"] == "tf":
        _add_docx_tf_options(doc, nq, student_mode=student_mode)
    elif nq["type"] in ("fill_in",):
        doc.add_paragraph("Answer: " + "_" * 40)
        # Word bank if available
        word_bank = nq.get("word_bank")
        if word_bank:
            p = doc.add_paragraph()
            run = p.add_run("Word Bank: " + ", ".join(word_bank))
            run.italic = True
            run.font.size = Pt(10)
    elif nq["type"] == "short_answer":
        if nq.get("rubric_hint") and not student_mode:
            p = doc.add_paragraph()
            run = p.add_run(f"(Hint: {nq['rubric_hint']})")
            run.italic = True
            run.font.size = Pt(9)
        doc.add_paragraph("_" * 60)
        doc.add_paragraph("_" * 60)
    elif nq["type"] == "ordering" and nq.get("ordering_items"):
        _add_docx_ordering(doc, nq)
    elif nq["type"] == "essay":
        for _ in range(4):
            doc.add_paragraph("_" * 60)
    elif nq["type"] == "matching" and nq["matches"]:
        _add_docx_matching(doc, nq)

    # Spacer
    doc.add_paragraph("")


def _add_docx_mc_options(doc, nq: dict, student_mode: bool = False):
    """Add multiple choice options with A/B/C/D lettering.

    In teacher mode the correct answer is bolded.
    In student mode all options render identically.
    """
    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for j, opt in enumerate(nq["options"]):
        letter = letters[j] if j < len(letters) else str(j)
        p = doc.add_paragraph()
        text = f"    {letter}. {opt}"
        run = p.add_run(text)
        if not student_mode and str(opt) == nq["correct_answer"]:
            run.bold = True


def _add_docx_tf_options(doc, nq: dict, student_mode: bool = False):
    """Add True/False options with A/B lettering.

    In teacher mode the correct answer is bolded.
    """
    for opt_letter, val in [("A", "True"), ("B", "False")]:
        p = doc.add_paragraph()
        run = p.add_run(f"    {opt_letter}. {val}")
        if not student_mode and val == nq["correct_answer"]:
            run.bold = True


def _add_docx_matching(doc, nq: dict):
    """Add a matching question as a simple table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Term"
    hdr[1].text = "Definition"
    for m in nq["matches"]:
        row = table.add_row().cells
        row[0].text = m["term"]
        row[1].text = m["definition"]


def _add_docx_ordering(doc, nq: dict):
    """Add an ordering question with numbered blanks for student response."""
    if nq.get("ordering_instructions"):
        p = doc.add_paragraph()
        run = p.add_run(nq["ordering_instructions"])
        run.italic = True
        run.font.size = Pt(9)

    # Show items in a scrambled display order (reverse of correct)
    items = nq.get("ordering_items", [])
    import random as _rng

    display_order = list(range(len(items)))
    # Use a deterministic shuffle based on question number
    _rng.Random(nq["number"]).shuffle(display_order)

    for rank, idx in enumerate(display_order):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"___  {items[idx]}")


def _add_docx_answer_key(doc, normalized: list):
    """Add answer key section listing all correct answers."""
    for nq in normalized:
        if nq["type"] == "matching" and nq["matches"]:
            answer_text = "; ".join(f"{m['term']} -> {m['definition']}" for m in nq["matches"])
        elif nq["type"] == "ordering" and nq.get("ordering_items"):
            correct_order = nq.get("ordering_correct_order", [])
            items = nq.get("ordering_items", [])
            ordered = []
            for idx in correct_order:
                if 0 <= idx < len(items):
                    ordered.append(items[idx])
            answer_text = " -> ".join(ordered) if ordered else "(see items)"
        elif nq["type"] == "short_answer":
            parts = [nq.get("expected_answer", "")]
            alts = nq.get("acceptable_answers", [])
            if alts:
                parts.append(f"(also: {', '.join(alts)})")
            answer_text = " ".join(parts)
        else:
            answer_text = nq["correct_answer"]

        p = doc.add_paragraph()
        run = p.add_run(f"{nq['number']}. ")
        run.bold = True
        p.add_run(answer_text or "(see rubric)")


# ---------------------------------------------------------------------------
# GIFT Export (Moodle / Canvas import format)
# ---------------------------------------------------------------------------


def export_gift(quiz, questions) -> str:
    """Export quiz to GIFT format for Moodle/Canvas import.

    Args:
        quiz: Quiz ORM object.
        questions: List of Question ORM objects.

    Returns:
        GIFT format string.
    """
    lines = []
    # Add quiz title as a comment
    lines.append(f"// {quiz.title or 'Quiz'}")
    lines.append("// Exported from QuizWeaver")
    lines.append("")

    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        gift_str = _format_gift_question(nq)
        lines.append(gift_str)
        lines.append("")

    return "\n".join(lines)


def _format_gift_question(nq: dict) -> str:
    """Format a single normalized question as GIFT."""
    title = f"Q{nq['number']}"
    escaped_text = _escape_gift(nq["text"])

    if nq["type"] == "mc":
        return _gift_mc(title, escaped_text, nq)
    elif nq["type"] == "tf":
        return _gift_tf(title, escaped_text, nq)
    elif nq["type"] == "short_answer":
        return _gift_short_answer(title, escaped_text, nq)
    elif nq["type"] == "fill_in":
        return _gift_short(title, escaped_text, nq)
    elif nq["type"] == "matching":
        return _gift_matching(title, escaped_text, nq)
    elif nq["type"] == "ordering":
        return _gift_ordering(title, escaped_text, nq)
    elif nq["type"] == "essay":
        return _gift_essay(title, escaped_text)
    else:
        # Default: treat as short answer
        return _gift_short(title, escaped_text, nq)


def _gift_mc(title: str, text: str, nq: dict) -> str:
    """Format MC question in GIFT."""
    parts = [f"::{title}:: {text} {{"]
    for opt in nq["options"]:
        escaped_opt = _escape_gift(str(opt))
        if str(opt) == nq["correct_answer"]:
            parts.append(f"  ={escaped_opt}")
        else:
            parts.append(f"  ~{escaped_opt}")
    parts.append("}")
    return "\n".join(parts)


def _gift_tf(title: str, text: str, nq: dict) -> str:
    """Format T/F question in GIFT."""
    answer = nq["correct_answer"].upper() if nq["correct_answer"] else "TRUE"
    return f"::{title}:: {text} {{{answer}}}"


def _gift_short(title: str, text: str, nq: dict) -> str:
    """Format short answer / fill-in question in GIFT."""
    answer = _escape_gift(nq["correct_answer"]) if nq["correct_answer"] else ""
    if answer:
        return f"::{title}:: {text} {{={answer}}}"
    return f"::{title}:: {text} {{}}"


def _gift_matching(title: str, text: str, nq: dict) -> str:
    """Format matching question in GIFT."""
    parts = [f"::{title}:: {text} {{"]
    for m in nq["matches"]:
        term = _escape_gift(m["term"])
        defn = _escape_gift(m["definition"])
        parts.append(f"  ={term} -> {defn}")
    parts.append("}")
    return "\n".join(parts)


def _gift_short_answer(title: str, text: str, nq: dict) -> str:
    """Format short answer question in GIFT with multiple acceptable answers."""
    answers = []
    expected = nq.get("expected_answer", "")
    if expected:
        answers.append(expected)
    for alt in nq.get("acceptable_answers", []):
        if alt and alt not in answers:
            answers.append(alt)
    if not answers and nq.get("correct_answer"):
        answers.append(nq["correct_answer"])

    if answers:
        answer_parts = " ".join(f"={_escape_gift(a)}" for a in answers)
        return f"::{title}:: {text} {{{answer_parts}}}"
    return f"::{title}:: {text} {{}}"


def _gift_ordering(title: str, text: str, nq: dict) -> str:
    """Format ordering question in GIFT.

    Moodle ordering format: each item gets a numbered position.
    Uses matching-style format where items map to their position number.
    """
    items = nq.get("ordering_items", [])
    correct_order = nq.get("ordering_correct_order", [])

    if not items:
        return f"::{title}:: {text} {{}}"

    # Build matching-style pairs: item -> position number
    parts = [f"::{title}:: {text} {{"]
    for pos, idx in enumerate(correct_order):
        if 0 <= idx < len(items):
            item = _escape_gift(items[idx])
            parts.append(f"  ={item} -> {pos + 1}")
    parts.append("}")
    return "\n".join(parts)


def _gift_essay(title: str, text: str) -> str:
    """Format essay question in GIFT."""
    return f"::{title}:: {text} {{}}"


# ---------------------------------------------------------------------------
# PDF Export
# ---------------------------------------------------------------------------


def export_pdf(
    quiz, questions, style_profile: Optional[dict] = None, student_mode: bool = False
) -> io.BytesIO:
    """Export quiz to a PDF document.

    Args:
        quiz: Quiz ORM object.
        questions: List of Question ORM objects.
        style_profile: Parsed style profile dict (optional).
        student_mode: If True, suppress correct-answer info inline,
            cognitive levels, image descriptions, and omit answer key.

    Returns:
        BytesIO buffer containing the PDF file.
    """
    if style_profile is None:
        style_profile = {}

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 50

    # --- Title page ---
    c.setFont("Helvetica-Bold", 18)
    title = quiz.title or "Quiz"
    c.drawCentredString(width / 2, y, title)
    y -= 30

    # Info block
    c.setFont("Helvetica", 10)
    info_lines = _pdf_info_lines(quiz, style_profile, student_mode=student_mode)
    for line in info_lines:
        c.drawCentredString(width / 2, y, line)
        y -= 14

    # Name / Date line for student copies
    if student_mode:
        y -= 10
        c.setFont("Helvetica", 11)
        c.drawString(50, y, "Name: ____________________________  Date: ____________")
        y -= 14

    y -= 20

    # --- Questions ---
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Questions")
    y -= 24

    normalized = []
    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        normalized.append(nq)
        y = _pdf_draw_question(c, nq, y, width, height, student_mode=student_mode)

    # --- Answer Key (new page) — teacher mode only ---
    if not student_mode:
        c.showPage()
        y = height - 50
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, "Answer Key")
        y -= 24

        c.setFont("Helvetica", 10)
        for nq in normalized:
            if y < 60:
                c.showPage()
                y = height - 50
            answer = _pdf_answer_text(nq)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, f"{nq['number']}.")
            c.setFont("Helvetica", 10)
            y = _pdf_draw_wrapped_text(c, answer, 72, y, width - 122, height)

    c.save()
    buf.seek(0)
    return buf


def _pdf_info_lines(quiz, style_profile: dict, student_mode: bool = False) -> List[str]:
    """Build info lines for the PDF title area."""
    lines = []
    sol = style_profile.get("sol_standards")
    if sol:
        if isinstance(sol, list):
            sol = ", ".join(sol)
        lines.append(f"Standards: {sol}")

    if not student_mode:
        framework = style_profile.get("cognitive_framework")
        if framework:
            lines.append(f"Framework: {framework.capitalize()}")

        difficulty = style_profile.get("difficulty")
        if difficulty:
            lines.append(f"Difficulty: {difficulty}/5")

        provider = style_profile.get("provider")
        model = style_profile.get("model")
        if provider:
            generated_by = model if model else provider
            lines.append(f"Generated by: {generated_by}")

    created = getattr(quiz, "created_at", None)
    if created:
        if isinstance(created, datetime):
            lines.append(f"Date: {created.strftime('%Y-%m-%d')}")
        else:
            lines.append(f"Date: {created}")

    return lines


def _pdf_draw_question(
    c, nq: dict, y: float, page_width: float, page_height: float, student_mode: bool = False
) -> float:
    """Draw a single question on the PDF canvas. Returns new y position."""
    # Check for page break
    if y < 120:
        c.showPage()
        y = page_height - 50

    # Header line: "1. [MC] (5 pts) - Remember"
    header_parts = [f"{nq['number']}."]
    header_parts.append(f"[{nq['type'].upper()}]")
    header_parts.append(f"({nq['points']} pts)")
    if nq["cognitive_level"] and not student_mode:
        header_parts.append(f"- {nq['cognitive_level']}")

    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, " ".join(header_parts))
    y -= 16

    # Question text with wrapping
    c.setFont("Helvetica", 10)
    y = _pdf_draw_wrapped_text(c, nq["text"], 50, y, page_width - 100, page_height)

    # Image description — teacher mode only
    if nq.get("image_description") and not student_mode:
        c.setFont("Helvetica-Oblique", 9)
        y = _pdf_draw_wrapped_text(c, f"[Image: {nq['image_description']}]", 60, y, page_width - 120, page_height)
        c.setFont("Helvetica", 10)

    # Options based on type
    if nq["type"] == "mc" and nq["options"]:
        letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        for j, opt in enumerate(nq["options"]):
            if y < 60:
                c.showPage()
                y = page_height - 50
            letter_label = letters[j] if j < len(letters) else str(j)
            c.drawString(70, y, f"{letter_label}. {opt}")
            y -= 14
    elif nq["type"] == "tf":
        for opt_letter, val in [("A", "True"), ("B", "False")]:
            if y < 60:
                c.showPage()
                y = page_height - 50
            c.drawString(70, y, f"{opt_letter}. {val}")
            y -= 14
    elif nq["type"] in ("fill_in",):
        # Wider blanks for fill-in questions
        c.drawString(70, y, "Answer: " + "_" * 40)
        y -= 14
        # Word bank if available
        word_bank = nq.get("word_bank")
        if word_bank:
            c.setFont("Helvetica-Oblique", 10)
            c.drawString(70, y, "Word Bank: " + ", ".join(word_bank))
            c.setFont("Helvetica", 10)
            y -= 14
    elif nq["type"] == "short_answer":
        if nq.get("rubric_hint") and not student_mode:
            c.setFont("Helvetica-Oblique", 9)
            c.drawString(70, y, f"(Hint: {nq['rubric_hint']})")
            y -= 14
            c.setFont("Helvetica", 10)
        c.drawString(70, y, "_" * 60)
        y -= 14
        c.drawString(70, y, "_" * 60)
        y -= 14
    elif nq["type"] == "ordering" and nq.get("ordering_items"):
        if nq.get("ordering_instructions"):
            c.setFont("Helvetica-Oblique", 9)
            c.drawString(70, y, nq["ordering_instructions"])
            y -= 14
            c.setFont("Helvetica", 10)
        items = nq["ordering_items"]
        import random as _rng

        display_order = list(range(len(items)))
        _rng.Random(nq["number"]).shuffle(display_order)
        for rank, idx in enumerate(display_order):
            if y < 60:
                c.showPage()
                y = page_height - 50
            c.drawString(70, y, f"___  {items[idx]}")
            y -= 14
    elif nq["type"] == "essay":
        for _ in range(4):
            if y < 60:
                c.showPage()
                y = page_height - 50
            c.drawString(70, y, "_" * 60)
            y -= 14
    elif nq["type"] == "matching" and nq["matches"]:
        # List terms with blanks
        for m in nq["matches"]:
            if y < 60:
                c.showPage()
                y = page_height - 50
            c.drawString(70, y, f"{m['term']}  ->  _______________")
            y -= 14
        # List definitions as a word bank (shuffled)
        import random as _rng_m
        definitions = [m["definition"] for m in nq["matches"]]
        _rng_m.Random(nq["number"]).shuffle(definitions)
        y -= 4
        if y < 60:
            c.showPage()
            y = page_height - 50
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(70, y, "Choices: " + "  |  ".join(definitions))
        c.setFont("Helvetica", 10)
        y -= 14

    y -= 10  # Spacer between questions
    return y


def _pdf_draw_wrapped_text(c, text: str, x: float, y: float, max_width: float, page_height: float) -> float:
    """Draw text with simple word wrapping. Returns new y position."""
    if not text:
        return y
    words = text.split()
    line = ""
    for word in words:
        test_line = f"{line} {word}".strip()
        if c.stringWidth(test_line, c._fontname, c._fontsize) < max_width:
            line = test_line
        else:
            if y < 60:
                c.showPage()
                y = page_height - 50
            c.drawString(x, y, line)
            y -= 14
            line = word
    if line:
        if y < 60:
            c.showPage()
            y = page_height - 50
        c.drawString(x, y, line)
        y -= 14
    return y


def _pdf_answer_text(nq: dict) -> str:
    """Get answer text for the answer key."""
    if nq["type"] == "matching" and nq["matches"]:
        return "; ".join(f"{m['term']} -> {m['definition']}" for m in nq["matches"])
    if nq["type"] == "ordering" and nq.get("ordering_items"):
        correct_order = nq.get("ordering_correct_order", [])
        items = nq.get("ordering_items", [])
        ordered = [items[idx] for idx in correct_order if 0 <= idx < len(items)]
        return " -> ".join(ordered) if ordered else "(see items)"
    if nq["type"] == "short_answer":
        parts = [nq.get("expected_answer", nq.get("correct_answer", ""))]
        alts = nq.get("acceptable_answers", [])
        if alts:
            parts.append(f"(also: {', '.join(alts)})")
        return " ".join(p for p in parts if p)
    return nq["correct_answer"] or "(see rubric)"


# ---------------------------------------------------------------------------
# QTI Export (Canvas-compatible)
# ---------------------------------------------------------------------------

# QTI 1.2 XML Templates

_QTI_MANIFEST_START = """<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="man_{manifest_id}" xmlns="http://www.imsglobal.org/xsd/imscp_v1p1">
  <metadata>
    <schema>IMS Content</schema>
    <schemaversion>1.1.3</schemaversion>
  </metadata>
  <organizations />
  <resources>
    <resource identifier="res_{assessment_id}" type="imsqti_xmlv1p2/imscc_xmlv1p1/assessment">
      <file href="{assessment_filename}"/>
    </resource>
  </resources>
</manifest>"""

_QTI_ASSESSMENT_HEADER = """<?xml version="1.0" encoding="UTF-8"?>
<questestinterop xmlns="http://www.imsglobal.org/xsd/ims_qtiasiv1p2">
  <assessment ident="{assessment_id}" title="{title}">
    <qtimetadata>
      <qtimetadatafield>
        <fieldlabel>qmd_assessmenttype</fieldlabel>
        <fieldentry>Examination</fieldentry>
      </qtimetadatafield>
    </qtimetadata>
    <section ident="root_section">
"""

_QTI_ASSESSMENT_FOOTER = """    </section>
  </assessment>
</questestinterop>"""


def _xml_escape(text: str) -> str:
    """Escape text for safe inclusion in XML."""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _qti_mc_item(ident: str, nq: dict) -> str:
    """Build a QTI MC item XML string from normalized question data."""
    text = _xml_escape(nq["text"])
    points = nq["points"] or 5

    # Find correct option index
    correct_idx = 0
    for i, opt in enumerate(nq["options"]):
        if str(opt) == nq["correct_answer"]:
            correct_idx = i
            break

    response_labels = ""
    for idx, opt in enumerate(nq["options"]):
        response_labels += (
            f'          <response_label ident="opt_{idx}">'
            f'<material><mattext texttype="text/plain">{_xml_escape(str(opt))}</mattext></material>'
            f"</response_label>\n"
        )

    return f"""      <item ident="{ident}" title="Q{nq["number"]}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>multiple_choice_question</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{text}</mattext></material>
          <response_lid ident="response1" rcardinality="Single">
            <render_choice>
{response_labels}            </render_choice>
          </response_lid>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><varequal respident="response1">opt_{correct_idx}</varequal></conditionvar>
            <setvar action="Set" varname="SCORE">{points}</setvar>
          </respcondition>
        </resprocessing>
      </item>"""


def _qti_tf_item(ident: str, nq: dict) -> str:
    """Build a QTI True/False item (implemented as MC with True/False options)."""
    is_true = nq["correct_answer"].strip().lower() == "true" if nq["correct_answer"] else True
    # Reuse MC with True/False options
    tf_nq = dict(nq)
    tf_nq["options"] = ["True", "False"]
    tf_nq["correct_answer"] = "True" if is_true else "False"
    return _qti_mc_item(ident, tf_nq)


def _qti_essay_item(ident: str, nq: dict) -> str:
    """Build a QTI essay item XML string."""
    text = _xml_escape(nq["text"])
    points = nq["points"] or 5

    return f"""      <item ident="{ident}" title="Q{nq["number"]}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>essay_question</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{text}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
          <respcondition continue="No">
            <conditionvar><other/></conditionvar>
          </respcondition>
        </resprocessing>
      </item>"""


def _qti_matching_item(ident: str, nq: dict) -> str:
    """Build a QTI matching item XML string."""
    text = _xml_escape(nq["text"])
    points = nq["points"] or 5
    matches = nq.get("matches", [])

    if not matches:
        return _qti_essay_item(ident, nq)

    # Build response groups - one per term, each with all definitions as choices
    response_grps = ""
    respconditions = ""
    for idx, match in enumerate(matches):
        term = _xml_escape(match.get("term", ""))

        # All definitions as response labels
        all_defs = ""
        for d_idx, m in enumerate(matches):
            d_text = _xml_escape(m.get("definition", ""))
            all_defs += (
                f'            <response_label ident="def_{d_idx}">'
                f'<material><mattext texttype="text/plain">{d_text}</mattext></material>'
                f'</response_label>\n'
            )

        response_grps += (
            f'        <response_grp ident="grp_{idx}">\n'
            f'          <material><mattext texttype="text/plain">{term}</mattext></material>\n'
            f'          <render_choice>\n{all_defs}'
            f'          </render_choice>\n'
            f'        </response_grp>\n'
        )

        respconditions += (
            f'          <respcondition>\n'
            f'            <conditionvar><varequal respident="grp_{idx}">def_{idx}</varequal></conditionvar>\n'
            f'            <setvar action="Add" varname="SCORE">{round(points / len(matches), 2)}</setvar>\n'
            f'          </respcondition>\n'
        )

    return f"""      <item ident="{ident}" title="Q{nq["number"]}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>matching_question</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{text}</mattext></material>
{response_grps}        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
{respconditions}        </resprocessing>
      </item>"""


def _qti_ordering_item(ident: str, nq: dict) -> str:
    """Build a QTI ordering item as matching (items to position numbers)."""
    items = nq.get("ordering_items", [])
    correct_order = nq.get("ordering_correct_order", [])

    if not items:
        return _qti_essay_item(ident, nq)

    # Convert to matching format: each item matches to its position
    matches = []
    for i, item in enumerate(items):
        pos = correct_order[i] + 1 if i < len(correct_order) else i + 1
        matches.append({"term": item, "definition": str(pos)})

    matching_nq = dict(nq)
    matching_nq["matches"] = matches
    return _qti_matching_item(ident, matching_nq)


def _qti_short_answer_item(ident: str, nq: dict) -> str:
    """Build a QTI short answer item with auto-grading support."""
    text = _xml_escape(nq["text"])
    points = nq["points"] or 5
    correct = _xml_escape(nq.get("correct_answer") or nq.get("expected_answer") or "")

    # Build varequal conditions for acceptable answers
    acceptable = nq.get("acceptable_answers", [])
    all_answers = [correct] + [_xml_escape(a) for a in acceptable if a]
    # Deduplicate while preserving order
    seen = set()
    unique_answers = []
    for a in all_answers:
        if a and a not in seen:
            seen.add(a)
            unique_answers.append(a)

    conditions = ""
    for ans in unique_answers:
        conditions += (
            f'          <respcondition continue="No">\n'
            f'            <conditionvar><varequal respident="response1" case="No">{ans}</varequal></conditionvar>\n'
            f'            <setvar action="Set" varname="SCORE">{points}</setvar>\n'
            f'          </respcondition>\n'
        )

    return f"""      <item ident="{ident}" title="Q{nq["number"]}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>short_answer_question</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material><mattext texttype="text/html">{text}</mattext></material>
          <response_str ident="response1" rcardinality="Single">
            <render_fib><response_label ident="answer1"/></render_fib>
          </response_str>
        </presentation>
        <resprocessing>
          <outcomes><decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/></outcomes>
{conditions}        </resprocessing>
      </item>"""


def export_qti(quiz, questions) -> io.BytesIO:
    """Export quiz as a QTI 1.2 ZIP package (Canvas-compatible).

    Args:
        quiz: Quiz ORM object.
        questions: List of Question ORM objects.

    Returns:
        BytesIO buffer containing the .zip file.
    """
    assessment_id = uuid.uuid4().hex
    manifest_id = uuid.uuid4().hex
    title = _xml_escape(quiz.title or "Quiz")
    assessment_filename = f"{assessment_id}.xml"

    # Build item XML for each question
    item_parts = []
    for i, q in enumerate(questions):
        nq = normalize_question(q, i)
        ident = f"q{nq['number']}"

        if nq["type"] == "mc":
            item_parts.append(_qti_mc_item(ident, nq))
        elif nq["type"] == "tf":
            item_parts.append(_qti_tf_item(ident, nq))
        elif nq["type"] == "matching" and nq.get("matches"):
            item_parts.append(_qti_matching_item(ident, nq))
        elif nq["type"] == "ordering" and nq.get("ordering_items"):
            item_parts.append(_qti_ordering_item(ident, nq))
        elif nq["type"] in ("short_answer", "fill_in"):
            item_parts.append(_qti_short_answer_item(ident, nq))
        elif nq["type"] == "essay":
            item_parts.append(_qti_essay_item(ident, nq))
        else:
            # Default: MC if options exist, else essay
            if nq["options"]:
                item_parts.append(_qti_mc_item(ident, nq))
            else:
                item_parts.append(_qti_essay_item(ident, nq))

    # Assemble assessment XML
    assessment_xml = _QTI_ASSESSMENT_HEADER.format(assessment_id=assessment_id, title=title)
    assessment_xml += "\n".join(item_parts)
    assessment_xml += "\n" + _QTI_ASSESSMENT_FOOTER

    # Assemble manifest XML
    manifest_xml = _QTI_MANIFEST_START.format(
        manifest_id=manifest_id,
        assessment_id=assessment_id,
        assessment_filename=assessment_filename,
    )

    # Write to ZIP in memory
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("imsmanifest.xml", manifest_xml)
        zf.writestr(assessment_filename, assessment_xml)

    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Legacy output functions (migrated from src/output.py)
#
# These are used by the old CLI pipeline in main.py and src/review.py.
# New code should prefer export_pdf() / export_qti() above.
# ---------------------------------------------------------------------------

# --- QTI XML Templates (legacy) ---

_MANIFEST_TEMPLATE_START = """<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="man_{manifest_id}" xmlns="http://www.imsglobal.org/xsd/imscp_v1p1" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xsi:schemaLocation="http://www.imsglobal.org/xsd/imscp_v1p1 http://www.imsglobal.org/xsd/imscp_v1p1.xsd">
  <metadata>
    <schema>IMS Content</schema>
    <schemaversion>1.1.3</schemaversion>
  </metadata>
  <organizations />
  <resources>
    <resource identifier="res_{assessment_id}" type="imsqti_xmlv1p2/imscc_xmlv1p1/assessment">
      <file href="{assessment_filename}"/>"""

_MANIFEST_TEMPLATE_END = """
    </resource>
  </resources>
</manifest>"""

_ASSESSMENT_HEADER = """<?xml version="1.0" encoding="UTF-8"?>
<questestinterop xmlns="http://www.imsglobal.org/xsd/ims_qtiasiv1p2" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xsi:schemaLocation="http://www.imsglobal.org/xsd/ims_qtiasiv1p2 http://www.imsglobal.org/xsd/ims_qtiasiv1p2.xsd">
  <assessment ident="{assessment_id}" title="{title}">
    <qtimetadata>
      <qtimetadatafield>
        <fieldlabel>qmd_assessmenttype</fieldlabel>
        <fieldentry>Examination</fieldentry>
      </qtimetadatafield>
    </qtimetadata>
    <section ident="root_section">
"""

_ASSESSMENT_FOOTER = """    </section>
  </assessment>
</questestinterop>"""


def _create_item_header(ident, title, points, image_ref=None, image_placeholder=None):
    img_tag = ""
    if image_ref:
        img_tag = f'<p><img src="{image_ref}" alt="Question Image"/></p>'
    elif image_placeholder:
        img_tag = f"<p><em>[{image_placeholder}]</em></p>"

    return f"""
      <item ident="{ident}" title="{title}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>{{qt}}</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material>
            <mattext texttype="text/html">{img_tag}{{question_text}}</mattext>
          </material>
          {{response_block}}
        </presentation>
        <resprocessing>
          <outcomes>
            <decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/>
          </outcomes>
          {{processing_block}}
        </resprocessing>
      </item>"""


def _create_mc_question_legacy(
    id,
    title,
    text,
    points,
    options,
    correct_index,
    image_ref=None,
    image_placeholder=None,
):
    response_labels = ""
    for idx, opt in enumerate(options):
        response_labels += f"""
          <response_label ident="opt_{idx}">
            <material><mattext texttype="text/plain">{opt}</mattext></material>
          </response_label>"""

    response_block = f"""<response_lid ident="response1" rcardinality="Single">
            <render_choice>
              {response_labels}
            </render_choice>
          </response_lid>"""

    processing_block = f"""<respcondition continue="No">
            <conditionvar>
              <varequal respident="response1">opt_{correct_index}</varequal>
            </conditionvar>
            <setvar action="Set" varname="SCORE">{points}</setvar>
          </respcondition>"""

    template = _create_item_header(id, title, points, image_ref, image_placeholder)
    return template.format(
        qt="multiple_choice_question",
        question_text=text,
        response_block=response_block,
        processing_block=processing_block,
    )


def _create_tf_question_legacy(id, title, text, points, is_true, image_ref=None, image_placeholder=None):
    return _create_mc_question_legacy(
        id,
        title,
        text,
        points,
        ["True", "False"],
        0 if is_true else 1,
        image_ref,
        image_placeholder,
    )


def _create_essay_question_legacy(id, title, text, points, image_ref=None, image_placeholder=None):
    response_block = """<response_str ident="response1" rcardinality="Single">
            <render_fib>
              <response_label ident="answer1"/>
            </render_fib>
          </response_str>"""

    processing_block = """<respcondition continue="No">
            <conditionvar>
              <other/>
            </conditionvar>
          </respcondition>"""

    template = _create_item_header(id, title, points, image_ref, image_placeholder)
    return template.format(
        qt="essay_question",
        question_text=text,
        response_block=response_block,
        processing_block=processing_block,
    )


def _create_multiple_answers_question_legacy(
    id,
    title,
    text,
    points,
    options,
    correct_indices,
    image_ref=None,
    image_placeholder=None,
):
    response_labels = ""
    for idx, opt in enumerate(options):
        response_labels += f"""
          <response_label ident="opt_{idx}">
            <material><mattext texttype="text/plain">{opt}</mattext></material>
          </response_label>"""

    response_block = f"""<response_lid ident="response1" rcardinality="Multiple">
            <render_choice>
              {response_labels}
            </render_choice>
          </response_lid>"""

    conditions = ""
    for i in range(len(options)):
        if i in correct_indices:
            conditions += f'<varequal respident="response1">opt_{i}</varequal>'
        else:
            conditions += f'<not><varequal respident="response1">opt_{i}</varequal></not>'

    processing_block = f"""<respcondition continue="No">
            <conditionvar>
              <and>
                {conditions}
              </and>
            </conditionvar>
            <setvar action="Set" varname="SCORE">{points}</setvar>
          </respcondition>"""

    template = _create_item_header(id, title, points, image_ref, image_placeholder)
    return template.format(
        qt="multiple_answers_question",
        question_text=text,
        response_block=response_block,
        processing_block=processing_block,
    )


def generate_pdf_preview(questions, filename, quiz_title, image_map=None):
    """Legacy PDF preview generator (from old output.py).

    Generates a simple PDF with questions and answer options.
    New code should use ``export_pdf()`` instead.
    """
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    y_position = height - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y_position, quiz_title)
    y_position -= 30

    c.setFont("Helvetica", 12)

    for i, q in enumerate(questions):
        if y_position < 100:
            c.showPage()
            y_position = height - 50

        q_text = q.get("text", "")
        q_type = q.get("type", "")
        q_title = q.get("title", f"Question {i + 1}")
        q_points = q.get("points", 5)
        image_ref = q.get("image_ref")
        image_placeholder = q.get("image_placeholder")

        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y_position, f"{i + 1}. {q_title} ({q_points} points)")
        y_position -= 20

        # Draw Image or Placeholder
        if image_ref and image_map and image_ref in image_map:
            try:
                img_path = image_map[image_ref]
                img = ImageReader(img_path)
                iw, ih = img.getSize()
                aspect = ih / float(iw)
                display_width = 400
                display_height = display_width * aspect

                if y_position - display_height < 50:
                    c.showPage()
                    y_position = height - 50

                c.drawImage(
                    img_path,
                    50,
                    y_position - display_height,
                    width=display_width,
                    height=display_height,
                )
                y_position -= display_height + 20
            except Exception as e:
                print(f"Error drawing image {image_ref}: {e}")

        elif image_placeholder:
            if y_position - 60 < 50:
                c.showPage()
                y_position = height - 50

            c.setStrokeColorRGB(0.5, 0.5, 0.5)
            c.rect(50, y_position - 50, 400, 50, fill=0)

            c.setFont("Helvetica-Oblique", 10)
            c.drawString(60, y_position - 20, "IMAGE PLACEHOLDER:")
            c.drawString(60, y_position - 35, image_placeholder[:80])
            y_position -= 70

        c.setFont("Helvetica", 10)

        text_object = c.beginText(50, y_position)
        text_object.setFont("Helvetica", 10)

        words = q_text.split()
        line = ""
        for word in words:
            if c.stringWidth(line + " " + word) < 500:
                line += " " + word
            else:
                text_object.textLine(line)
                y_position -= 12
                line = word
        text_object.textLine(line)
        y_position -= 12
        c.drawText(text_object)
        y_position -= 10

        if q_type == "mc":
            options = q.get("options", [])
            correct_idx = q.get("correct_index", -1)
            for idx, opt in enumerate(options):
                prefix = "[x]" if idx == correct_idx else "[ ]"
                c.drawString(70, y_position, f"{prefix} {opt}")
                y_position -= 15
        elif q_type == "tf":
            is_true = q.get("is_true", False)
            c.drawString(70, y_position, "[x] True" if is_true else "[ ] True")
            y_position -= 15
            c.drawString(70, y_position, "[ ] False" if is_true else "[x] False")
            y_position -= 15

        y_position -= 20

    c.save()


def create_qti_package(questions, used_images, config):
    """Legacy QTI package generator (from old output.py).

    Generates a QTI ZIP package for Canvas import.
    New code should use ``export_qti()`` instead.
    """
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    assessment_id = uuid.uuid4().hex
    manifest_id = uuid.uuid4().hex

    question_xml_parts = []
    for i, q in enumerate(questions):
        q_id = f"q{i + 1}"
        image_ref = q.get("image_ref")
        image_placeholder = q.get("image_placeholder")
        q_type = q.get("type", "mc")

        if q_type == "mc":
            question_xml_parts.append(
                _create_mc_question_legacy(
                    q_id, q.get("title", "Question"), q.get("text", ""),
                    q.get("points", 5), q.get("options", []),
                    q.get("correct_index", 0), image_ref, image_placeholder,
                )
            )
        elif q_type == "tf":
            question_xml_parts.append(
                _create_tf_question_legacy(
                    q_id, q.get("title", "Question"), q.get("text", ""),
                    q.get("points", 5), q.get("is_true", True),
                    image_ref, image_placeholder,
                )
            )
        elif q_type == "essay":
            question_xml_parts.append(
                _create_essay_question_legacy(
                    q_id, q.get("title", "Question"), q.get("text", ""),
                    q.get("points", 5), image_ref, image_placeholder,
                )
            )
        elif q_type == "ma":
            correct_indices = q.get("correct_indices", [])
            if not correct_indices and "correct_index" in q:
                correct_indices = [q["correct_index"]]
            question_xml_parts.append(
                _create_multiple_answers_question_legacy(
                    q_id, q.get("title", "Question"), q.get("text", ""),
                    q.get("points", 5), q.get("options", []),
                    correct_indices, image_ref, image_placeholder,
                )
            )
        else:
            print(f"Warning: Unknown question type {q_type}, defaulting to MC")
            question_xml_parts.append(
                _create_mc_question_legacy(
                    q_id, q.get("title", "Question"), q.get("text", ""),
                    q.get("points", 5), q.get("options", []),
                    q.get("correct_index", 0), image_ref, image_placeholder,
                )
            )

    quiz_content = _ASSESSMENT_HEADER.format(
        assessment_id=assessment_id, title=config["generation"]["quiz_title"]
    )
    quiz_content += "\n".join(question_xml_parts)
    quiz_content += _ASSESSMENT_FOOTER

    manifest_content = _MANIFEST_TEMPLATE_START.format(
        manifest_id=manifest_id,
        assessment_id=assessment_id,
        assessment_filename=config["qti"]["assessment_filename"],
    )
    for _, img_name in used_images:
        manifest_content += f'\n      <file href="{img_name}"/>'
    manifest_content += _MANIFEST_TEMPLATE_END

    zip_filename = os.path.join(
        config["paths"]["quiz_output_dir"],
        config["qti"]["zip_filename_template"].format(timestamp=timestamp),
    )

    with zipfile.ZipFile(zip_filename, "w") as zipf:
        zipf.writestr(config["qti"]["manifest_name"], manifest_content)
        zipf.writestr(config["qti"]["assessment_filename"], quiz_content)
        for img_path, img_name in used_images:
            zipf.write(img_path, img_name)

    return zip_filename
